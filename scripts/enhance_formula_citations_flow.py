from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


U_REF_HEADING = "\u53c2\u8003\u6587\u732e"
U_APPENDIX = "\u9644\u5f55"
U_TECH_ROUTE_HEADING = "1.5 \u6280\u672f\u8def\u7ebf\u4e0e\u53ef\u884c\u6027\u5206\u6790"
U_TECH_ROUTE_CAPTION = "\u56fe 1.1 \u7cfb\u7edf\u6280\u672f\u8def\u7ebf\u56fe"
U_TECH_ROUTE_SENTENCE = "\u6574\u4f53\u6280\u672f\u8def\u7ebf\u5982\u56fe 1.1 \u6240\u793a\u3002"


def normalize(text: str) -> str:
    return " ".join(text.split())


def set_run_font(run, east="\u5b8b\u4f53", west="Times New Roman", size=Pt(12), superscript=None, color=None, italic=None):
    run.font.name = west
    run.font.size = size
    if italic is not None:
        run.italic = italic
    if superscript is not None:
        run.font.superscript = superscript
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:cs"), west)


def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def find_heading(document: Document, text: str) -> int:
    for i, paragraph in enumerate(document.paragraphs):
        if normalize(paragraph.text) == text:
            return i
    raise ValueError(f"Heading not found: {text}")


def find_references(document: Document) -> tuple[int, int, list[tuple[int, int]]]:
    texts = [normalize(p.text) for p in document.paragraphs]
    ref_idx = next((i for i, t in enumerate(texts) if t == U_REF_HEADING), -1)
    if ref_idx < 0:
        raise ValueError("Reference heading not found")
    app_idx = next((i for i, t in enumerate(texts) if i > ref_idx and t.startswith(U_APPENDIX)), len(texts))
    refs: list[tuple[int, int]] = []
    for i in range(ref_idx + 1, app_idx):
        match = re.match(r"^\[(\d+)]", texts[i])
        if match:
            refs.append((int(match.group(1)), i))
    return ref_idx, app_idx, refs


def max_bookmark_id(document: Document) -> int:
    max_id = 0
    for paragraph in document.paragraphs:
        for bookmark in paragraph._element.xpath(".//w:bookmarkStart"):
            raw = bookmark.get(qn("w:id"))
            if raw and raw.isdigit():
                max_id = max(max_id, int(raw))
    return max_id


def existing_bookmarks(document: Document) -> set[str]:
    names = set()
    for paragraph in document.paragraphs:
        for bookmark in paragraph._element.xpath(".//w:bookmarkStart"):
            name = bookmark.get(qn("w:name"))
            if name:
                names.add(name)
    return names


def add_reference_bookmarks(document: Document) -> dict[int, str]:
    _, _, refs = find_references(document)
    names = existing_bookmarks(document)
    next_id = max_bookmark_id(document) + 1
    ref_anchors: dict[int, str] = {}
    for ref_no, p_idx in refs:
        anchor = f"ref_{ref_no}"
        ref_anchors[ref_no] = anchor
        if anchor in names:
            continue
        paragraph = document.paragraphs[p_idx]
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(next_id))
        start.set(qn("w:name"), anchor)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(next_id))
        insert_at = 1 if paragraph._p.pPr is not None else 0
        paragraph._p.insert(insert_at, start)
        paragraph._p.append(end)
        next_id += 1
    return ref_anchors


def make_text_run(text: str, superscript: bool) -> OxmlElement:
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    r_pr.append(size)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "24")
    r_pr.append(size_cs)
    if superscript:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        r_pr.append(va)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    return run


def make_hyperlink_run(text: str, anchor: str, superscript: bool) -> OxmlElement:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = make_text_run(text, superscript)
    r_pr = run.get_or_add_rPr()
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)
    hyperlink.append(run)
    return hyperlink


def add_citation_hyperlinks(document: Document, ref_anchors: dict[int, str]) -> int:
    ref_idx, _, _ = find_references(document)
    changed = 0
    marker_re = re.compile(r"^\[([0-9,\-\s]+)]$")
    for p_idx, paragraph in enumerate(document.paragraphs[:ref_idx]):
        for run in list(paragraph.runs):
            text = run.text.strip()
            if not run.font.superscript:
                continue
            match = marker_re.match(text)
            if not match:
                continue
            content = match.group(1).replace(" ", "")
            parent = run._element.getparent()
            insert_pos = parent.index(run._element)
            parent.remove(run._element)
            new_nodes: list[OxmlElement] = [make_text_run("[", True)]
            pos = 0
            for token in re.finditer(r"\d+", content):
                if token.start() > pos:
                    new_nodes.append(make_text_run(content[pos : token.start()], True))
                ref_no = int(token.group(0))
                if ref_no in ref_anchors:
                    new_nodes.append(make_hyperlink_run(token.group(0), ref_anchors[ref_no], True))
                else:
                    new_nodes.append(make_text_run(token.group(0), True))
                pos = token.end()
            if pos < len(content):
                new_nodes.append(make_text_run(content[pos:], True))
            new_nodes.append(make_text_run("]", True))
            for offset, node in enumerate(new_nodes):
                parent.insert(insert_pos + offset, node)
            changed += 1
    return changed


def current_chapter_by_index(document: Document, p_idx: int) -> int:
    chapter = 1
    chinese = {
        "\u4e00": 1,
        "\u4e8c": 2,
        "\u4e09": 3,
        "\u56db": 4,
        "\u4e94": 5,
        "\u516d": 6,
        "\u4e03": 7,
        "\u516b": 8,
        "\u4e5d": 9,
    }
    for paragraph in document.paragraphs[: p_idx + 1]:
        text = normalize(paragraph.text)
        match = re.match(r"^\u7b2c([%s])\u7ae0" % "".join(chinese.keys()), text)
        if match:
            chapter = chinese[match.group(1)]
    return chapter


def clear_paragraph_keep_ppr(paragraph: Paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def append_tab_run(paragraph: Paragraph):
    run = OxmlElement("w:r")
    tab = OxmlElement("w:tab")
    run.append(tab)
    paragraph._p.append(run)


def append_number_run(paragraph: Paragraph, number: str):
    run = paragraph.add_run()
    run.add_tab()
    run.add_text(number)
    set_run_font(run, "\u5b8b\u4f53", "Times New Roman", Pt(12))


def append_formula_text_run(paragraph: Paragraph, formula: str):
    run = paragraph.add_run()
    run.add_text(formula)
    set_run_font(run, "\u5b8b\u4f53", "Times New Roman", Pt(12), italic=True)


def rendered_formula_override(math_text: str) -> str | None:
    compact = " ".join(math_text.split())
    if compact.startswith("P(no collision)"):
        return "P(no collision) = \u220f\u1d62\u208c\u2080\u1d4f\u207b\u00b9 (1 - i / 2\u00b2\u2075\u2076)"
    if compact.startswith("P(collision) = 1"):
        return "P(collision) = 1 - \u220f\u1d62\u208c\u2080\u1d4f\u207b\u00b9 (1 - i / 2\u00b2\u2075\u2076)"
    if compact.startswith("k <<"):
        return "k \u226a 2\u00b9\u00b2\u2078, P(collision) \u2248 k(k - 1) / 2\u00b2\u2075\u2077"
    return None


def format_formulas(document: Document) -> int:
    formula_indices = []
    for i, paragraph in enumerate(document.paragraphs):
        xml = paragraph._element.xml
        text = normalize(paragraph.text)
        if ("<m:oMath" in xml or "<m:oMathPara" in xml) and not text:
            formula_indices.append(i)

    counters: dict[int, int] = {}
    for p_idx in formula_indices:
        paragraph = document.paragraphs[p_idx]
        original_math_text = " ".join((node.text or "") for node in paragraph._element.xpath(".//m:t"))
        plain_override = rendered_formula_override(original_math_text)
        math_children = [copy.deepcopy(child) for child in paragraph._p if child.tag in {qn("m:oMath"), qn("m:oMathPara")}]
        if not math_children:
            continue
        chapter = current_chapter_by_index(document, p_idx)
        counters[chapter] = counters.get(chapter, 0) + 1
        number = f"({chapter}.{counters[chapter]})"

        clear_paragraph_keep_ppr(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(7.8), WD_TAB_ALIGNMENT.CENTER)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.6), WD_TAB_ALIGNMENT.RIGHT)

        append_tab_run(paragraph)
        if plain_override:
            append_formula_text_run(paragraph, plain_override)
        else:
            for math in math_children:
                paragraph._p.append(math)
        append_number_run(paragraph, number)
    return len(formula_indices)


def font_path() -> str | None:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    return None


def generate_flowchart(path: Path):
    width, height = 1800, 980
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    fpath = font_path()
    title_font = ImageFont.truetype(fpath, 50) if fpath else ImageFont.load_default()
    node_font = ImageFont.truetype(fpath, 40) if fpath else ImageFont.load_default()
    small_font = ImageFont.truetype(fpath, 30) if fpath else ImageFont.load_default()

    draw.text((width // 2, 52), "\u53bb\u4e2d\u5fc3\u5316\u6295\u7968\u7cfb\u7edf\u6280\u672f\u8def\u7ebf", fill="#1D3557", font=title_font, anchor="mm")
    nodes = [
        ("\u9700\u6c42\u5206\u6790", "\u89d2\u8272\u3001\u7ea6\u675f\u3001\u5a01\u80c1\u8fb9\u754c", "#D8EAF5"),
        ("\u5408\u7ea6\u5efa\u6a21", "VotingSystem\u3001MerkleRoot", "#FDE7C7"),
        ("\u672c\u5730\u6d4b\u8bd5", "Hardhat\u3001Gas\u3001\u8fb9\u754c\u7528\u4f8b", "#E4F2D7"),
        ("\u6d4b\u8bd5\u7f51\u90e8\u7f72", "Sepolia\u3001\u5730\u5740\u3001\u4ea4\u6613 Hash", "#E8E0F4"),
        ("\u524d\u7aef\u8054\u8c03", "React\u3001Ethers.js\u3001MetaMask", "#DDECFB"),
        ("\u8bc1\u636e\u5f52\u6863", "\u622a\u56fe\u3001\u5ba1\u8ba1\u3001\u53ef\u590d\u6838\u62a5\u544a", "#E7E7E7"),
    ]
    positions = [
        (110, 220),
        (650, 220),
        (1190, 220),
        (1190, 525),
        (650, 525),
        (110, 525),
    ]
    node_w, node_h = 430, 205
    border = "#457B9D"
    for idx, (title, desc, fill) in enumerate(nodes):
        x, y = positions[idx]
        draw.rounded_rectangle((x, y, x + node_w, y + node_h), radius=24, fill=fill, outline=border, width=4)
        draw.text((x + node_w / 2, y + 55), title, fill="#1D3557", font=node_font, anchor="mm")
        draw.text((x + node_w / 2, y + 128), desc, fill="#333333", font=small_font, anchor="mm")

    arrow_pairs = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5)]
    for a, b in arrow_pairs:
        x1, y1 = positions[a]
        x2, y2 = positions[b]
        if a in {0, 1}:
            start = (x1 + node_w + 14, y1 + node_h / 2)
            end = (x2 - 14, y2 + node_h / 2)
            head = [(end[0], end[1]), (end[0] - 22, end[1] - 13), (end[0] - 22, end[1] + 13)]
        elif a == 2:
            start = (x1 + node_w / 2, y1 + node_h + 14)
            end = (x2 + node_w / 2, y2 - 14)
            head = [(end[0], end[1]), (end[0] - 13, end[1] - 22), (end[0] + 13, end[1] - 22)]
        else:
            start = (x1 - 14, y1 + node_h / 2)
            end = (x2 + node_w + 14, y2 + node_h / 2)
            head = [(end[0], end[1]), (end[0] + 22, end[1] - 13), (end[0] + 22, end[1] + 13)]
        draw.line((start[0], start[1], end[0], end[1]), fill="#1D3557", width=6)
        draw.polygon(head, fill="#1D3557")

    draw.rounded_rectangle((220, 820, 1580, 915), radius=18, fill="#F7FAFC", outline="#A8DADC", width=3)
    draw.text(
        (900, 868),
        "\u4ee5\u201c\u94fe\u4e0a\u5408\u7ea6\u7ea6\u675f + \u94fe\u4e0b\u8bc1\u636e\u5f52\u6863\u201d\u5f62\u6210\u53ef\u8fd0\u884c\u3001\u53ef\u6d4b\u8bd5\u3001\u53ef\u590d\u6838\u7684\u5de5\u7a0b\u95ed\u73af",
        fill="#1D3557",
        font=small_font,
        anchor="mm",
    )
    image.save(path, dpi=(450, 450))


def add_figure_bookmark(paragraph: Paragraph, name: str):
    names = {b.get(qn("w:name")) for b in paragraph._element.xpath(".//w:bookmarkStart")}
    if name in names:
        return
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "9001")
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "9001")
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_tech_route_figure(document: Document, image_path: Path) -> bool:
    if any(U_TECH_ROUTE_CAPTION in paragraph.text for paragraph in document.paragraphs):
        return False
    heading_idx = find_heading(document, U_TECH_ROUTE_HEADING)
    anchor = None
    for paragraph in document.paragraphs[heading_idx + 1 :]:
        if normalize(paragraph.text).startswith("1.") or normalize(paragraph.text).startswith("\u7b2c"):
            break
        if paragraph.text.strip():
            anchor = paragraph
            break
    if anchor is None:
        raise ValueError("Could not find 1.5 body paragraph")
    if U_TECH_ROUTE_SENTENCE not in anchor.text:
        run = anchor.add_run(U_TECH_ROUTE_SENTENCE)
        set_run_font(run, "\u5b8b\u4f53", "Times New Roman", Pt(12))

    img_para = insert_paragraph_after(anchor)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(6)
    img_para.paragraph_format.space_after = Pt(0)
    img_run = img_para.add_run()
    img_run.add_picture(str(image_path), width=Cm(14.5))

    caption = insert_paragraph_after(img_para)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.25
    run = caption.add_run(U_TECH_ROUTE_CAPTION)
    set_run_font(run, "\u5b8b\u4f53", "Times New Roman", Pt(10.5))
    add_figure_bookmark(caption, "fig_1_1")
    return True


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    parser.add_argument("--asset-dir", default="artifacts/generated_assets")
    args = parser.parse_args()

    asset_dir = Path(args.asset_dir)
    asset_dir.mkdir(parents=True, exist_ok=True)
    flowchart = asset_dir / "tech_route_flowchart.png"
    generate_flowchart(flowchart)

    document = Document(args.input)
    ref_anchors = add_reference_bookmarks(document)
    citation_count = add_citation_hyperlinks(document, ref_anchors)
    formula_count = format_formulas(document)
    figure_added = add_tech_route_figure(document, flowchart)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    print(
        {
            "output": str(output),
            "reference_bookmarks": len(ref_anchors),
            "citation_markers_hyperlinked": citation_count,
            "formula_paragraphs_formatted": formula_count,
            "flowchart_added": figure_added,
            "flowchart": str(flowchart),
        }
    )


if __name__ == "__main__":
    main()
