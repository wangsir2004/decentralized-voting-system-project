from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


BODY_11 = [
    "电子投票是组织决策、校园治理和在线协作中常见的信息化应用。与纸质投票相比，电子投票能够降低人工计票成本，提高结果统计效率，并扩大远程参与范围；但传统电子投票系统通常依赖中心化服务器和数据库保存选民名单、投票状态与计票结果。一旦系统运营方权限过大、数据库日志不完整或服务器遭受攻击，投票过程就容易出现难以独立复核的问题。已有研究将区块链电子投票视为提升过程透明度和结果可验证性的重要方向[[1]]。",
    "区块链通过分布式账本、哈希链接和共识机制记录交易，能够在不依赖单一中心机构的前提下保存可追溯数据[[2-3]]。在 Ethereum 平台中，智能合约可以把投票资格校验、重复投票限制、投票窗口控制和计票逻辑写入链上程序，使投票规则从平台承诺转化为可检查、可执行的代码约束[[4-6]]。因此，将区块链和智能合约引入电子投票系统，不是单纯改变数据存储位置，而是尝试把投票过程中的关键约束公开化、自动化和可审计化。",
    "本文课题围绕“基于 Solidity 智能合约与 Ethereum 测试网的去中心化电子投票系统设计与实现”展开，研究目标并不是构建可直接用于真实政务选举的生产级系统，而是完成一个面向本科毕业设计和教学演示的小规模可运行原型。系统重点验证智能合约投票、Merkle Tree 白名单、钱包签名、Sepolia 测试网部署、链上结果查询和安全审计等技术路径，并通过合约地址、交易 Hash、Gas 数据和运行截图保留可复核证据。相关以太坊投票系统研究和平台实现工作表明，这类原型适合用于分析去中心化投票的可行性、边界和工程实现过程[[7-8]]。",
]


BODY_12 = [
    "国内外电子投票研究长期围绕身份认证、投票匿名性、公开验证、抗篡改和系统可用性展开。传统电子投票方案多依赖中心化身份认证、数据库事务和审计日志来保障流程完整性，工程实现较成熟，但可信基础仍集中在系统运营方和数据库管理方。区块链电子投票研究则试图利用不可篡改账本、智能合约和密码学证明增强过程可验证性；近年系统综述也显示，区块链电子投票已经形成智能合约自动化、隐私保护和分布式验证等多条研究路线[[1,9]]。",
    "国内研究较早从去中心化架构和智能合约投票协议入手。基于以太坊智能合约的投票系统设计研究，强调将投票创建、资格校验和计票逻辑写入链上合约，以降低中心化服务器对投票结果的控制力[[7]]。吕佳卓围绕智能合约去中心化投票协议展开研究，将可信公告板、可链接环签名和门限加密机制结合起来，以增强投票结果的可验证性[[10]]；董友康则将盲签名算法与区块链平台结合，通过智能合约管理投票流程和计票过程，突出系统公开性与可信度[[11]]。这些工作为本文采用 Solidity 合约实现投票规则提供了直接的研究基础。",
    "随着研究深入，国内学者逐渐把重点扩展到公平性、隐私保护和复杂密码机制。周敏针对自计票电子投票中可能提前获知结果的问题，引入区块链、同态加密和代理签名改进投票流程[[12]]；吴佳龙通过可链接环签名保护选民匿名性并保持投票过程可审计[[13]]；戴鑫淦结合智能合约、可链接环签名和阈值加密降低管理员控制权[[14]]；田嗣犇则从秘密共享、同态加密和盲签名等角度研究投票区块链的隐私保护[[15]]。这些研究说明，强匿名和强隐私投票通常需要多种密码学机制协同完成，系统复杂度明显高于普通工程演示原型。",
    "国外研究同样关注智能合约投票、平台实现和系统性评估。McCorry 等较早讨论了利用智能合约实现组织投票并兼顾隐私保护的问题[[16]]；Daraghmi 等提出 VoteChain 架构，将投票规则和验证逻辑写入智能合约，以支持投票过程透明和结果可追溯[[17]]；Abdul 等基于 Ethereum、智能合约和 MetaMask 钱包实现去中心化电子投票系统，体现了区块链平台、钱包交互和前端应用结合的工程路径[[8]]。Sánchez 等对 2022—2025 年区块链电子投票研究进行系统评估，指出相关研究已经覆盖智能合约自动化、先进密码学机制和分布式账本验证等方向，但在系统性能、可扩展性和实际部署验证方面仍存在差异[[9]]。",
    "综合来看，已有研究为区块链电子投票提供了较为充分的理论和方案基础，但不同研究目标之间存在明显侧重：隐私保护类方案强调匿名性与抗关联分析，通常需要环签名、同态加密、零知识证明或门限加密等复杂机制；工程实现类方案更关注合约规则、钱包签名、链上计票和部署证据。本文选择后一路径，将系统定位为小规模可验证原型，重点完成白名单资格控制、防重复投票、链上公开计票、测试网部署和安全审计。这样的定位能够在本科毕业设计周期内形成可运行、可测试、可复核的工程闭环，同时也保留对真实身份认证、强匿名投票和大规模并发等问题的边界说明[[8,15]]。",
]


P13_TARGET_PREFIX = "围绕上述目标，本文完成以下工作："
P13_REPLACEMENT = (
    "围绕上述目标，本文完成以下工作：第一，设计投票系统的角色、功能需求、非功能需求和威胁边界；第二，基于 Solidity 编写 VotingSystem 智能合约，并使用 OpenZeppelin MerkleProof 完成白名单验证[[18-19]]；第三，使用 Hardhat 搭建编译、测试、部署和 Gas 统计流程[[20]]；第四，使用 React、TypeScript、Ethers.js 和 MetaMask 实现前端 DApp[[21]]；第五，在 Sepolia 测试网上重新部署合约并提交真实投票交易；第六，使用 Slither 完成合约静态分析并形成审计记录[[22]]。"
)

P15_TARGET_PREFIX = "从技术可行性看，Ethereum 账户模型能够天然区分不同投票地址"
P15_REPLACEMENT = (
    "从技术可行性看，Ethereum 账户模型能够天然区分不同投票地址，智能合约能够在链上保存不可抵赖的状态更新记录[[4-5]]；ECDSA 签名机制为账户交易提供身份校验基础[[23]]，Merkle Tree 能够用一个根哈希表示完整白名单并支持对数级 proof 验证[[24]]，React 前端能够通过钱包插件完成用户交互。因此，课题所需的关键技术均有成熟生态支撑。项目没有设计自定义密码学算法，而是使用 Solidity 语言特性、OpenZeppelin 安全库和 Ethers.js 通用接口完成实现，这有助于降低毕业设计阶段的实现风险。"
)


REFERENCES = [
    "蔡维德, 郁莲, 王星, 等. 基于区块链的电子投票系统研究[J]. 软件学报, 2018, 29(10): 2912-2933.",
    "袁勇, 王飞跃. 区块链技术发展现状与展望[J]. 自动化学报, 2016, 42(4): 481-494.",
    "Nakamoto S. Bitcoin: A Peer-to-Peer Electronic Cash System[EB/OL]. 2008.",
    "Buterin V. Ethereum Whitepaper: A Next-Generation Smart Contract and Decentralized Application Platform[EB/OL]. 2014.",
    "Wood G. Ethereum: A Secure Decentralised Generalised Transaction Ledger[R]. Ethereum Yellow Paper, 2024.",
    "郑子彬, 麦恩明, 连晓聪, 等. 区块链智能合约研究进展[J]. 自动化学报, 2022, 48(11): 2641-2656.",
    "王继成, 胡建军, 张建标. 基于以太坊智能合约的去中心化投票系统设计[J]. 计算机工程与科学, 2020, 42(7): 1195-1201.",
    "Abdul S A, Sarthak S, Jitender S, et al. Decentralized E-Voting System Using Ethereum Blockchain Technology[J]. Advances in Science and Technology, 2023, 124: 619-627. DOI:10.4028/P-0ZWK07.",
    "Sánchez R O, Salazar B A, González B M. Democratic Innovation: Systematic Evaluation of Blockchain-Based Electronic Voting (2022-2025)[J]. Technologies, 2026, 14(2): 95. DOI:10.3390/TECHNOLOGIES14020095.",
    "吕佳卓. 基于智能合约的去中心化安全电子投票系统[D]. 哈尔滨工业大学, 2019. DOI:10.27061/d.cnki.ghgdu.2019.004687.",
    "董友康. 基于区块链的安全电子投票系统的设计与实现[D]. 北京交通大学, 2019.",
    "周敏. 基于区块链的安全电子投票方案研究[D]. 南京邮电大学, 2022. DOI:10.27251/d.cnki.gnjdc.2022.000507.",
    "吴佳龙. 基于区块链和可链接环签名的电子投票方案研究[D]. 西安电子科技大学, 2023. DOI:10.27389/d.cnki.gxadu.2023.003086.",
    "戴鑫淦. 基于智能合约的隐私保护电子投票技术的研究[D]. 暨南大学, 2023. DOI:10.27167/d.cnki.gjinu.2023.000850.",
    "田嗣犇. 投票区块链的隐私保护技术研究[D]. 昆明理工大学, 2024. DOI:10.27200/d.cnki.gkmlu.2024.002020.",
    "McCorry P, Shahandashti S F, Hao F. A smart contract for boardroom voting with maximum voter privacy[C]//Financial Cryptography and Data Security. Cham: Springer, 2017: 357-375.",
    "Daraghmi E, Hamoudi A, Helou A M. Decentralizing Democracy: Secure and Transparent E-Voting Systems with Blockchain Technology in the Context of Palestine[J]. Future Internet, 2024, 16(11): 388. DOI:10.3390/FI16110388.",
    "Solidity Team. Solidity Documentation, Version 0.8.x[EB/OL].",
    "OpenZeppelin. Contracts Documentation: MerkleProof[EB/OL].",
    "Hardhat. Ethereum Development Environment Documentation[EB/OL].",
    "Ethers.js Contributors. Ethers.js Documentation[EB/OL].",
    "Trail of Bits. Slither: Static Analyzer for Solidity[EB/OL].",
    "Johnson D, Menezes A, Vanstone S. The Elliptic Curve Digital Signature Algorithm (ECDSA)[J]. International Journal of Information Security, 2001, 1(1): 36-63.",
    "Merkle R C. A Digital Signature Based on a Conventional Encryption Function[C]//CRYPTO. Berlin: Springer, 1987: 369-378.",
]


def normalize(text: str) -> str:
    return " ".join(text.split())


def set_run_font(run, east="宋体", west="Times New Roman", size=Pt(12), bold=None, superscript=None):
    run.font.name = west
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if superscript is not None:
        run.font.superscript = superscript
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:cs"), west)


def clear_paragraph(paragraph: Paragraph):
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def add_text_with_citations(paragraph: Paragraph, text: str, size=Pt(12), east="宋体", west="Times New Roman"):
    clear_paragraph(paragraph)
    parts = re.split(r"(\[\[[0-9,\-\s]+]])", text)
    for part in parts:
        if not part:
            continue
        citation = re.fullmatch(r"\[\[([0-9,\-\s]+)]]", part)
        if citation:
            run = paragraph.add_run(f"[{citation.group(1).replace(' ', '')}]")
            set_run_font(run, east=east, west=west, size=size, superscript=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, east=east, west=west, size=size)


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    add_text_with_citations(new_para, text)
    format_body(new_para)
    return new_para


def format_body(paragraph: Paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def format_reference(paragraph: Paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, "宋体", "Times New Roman", Pt(10.5))


def find_paragraph(document: Document, exact: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if normalize(paragraph.text) == exact:
            return paragraph
    raise ValueError(f"Paragraph not found: {exact}")


def remove_between(start: Paragraph, end: Paragraph):
    body = start._element.getparent()
    current = start._element.getnext()
    while current is not None and current is not end._element:
        nxt = current.getnext()
        body.remove(current)
        current = nxt


def insert_many_after(anchor: Paragraph, texts: list[str]):
    current = anchor
    for text in texts:
        current = insert_paragraph_after(current, text)


def replace_matching_paragraph(document: Document, prefix: str, replacement: str):
    for paragraph in document.paragraphs:
        if normalize(paragraph.text).startswith(prefix):
            add_text_with_citations(paragraph, replacement)
            format_body(paragraph)
            return
    raise ValueError(f"Target paragraph not found: {prefix}")


def replace_references(document: Document):
    ref_heading = find_paragraph(document, "参考文献")
    appendix_heading = None
    seen_ref_heading = False
    for paragraph in document.paragraphs:
        if paragraph._element is ref_heading._element:
            seen_ref_heading = True
            continue
        if not seen_ref_heading:
            continue
        if normalize(paragraph.text).startswith("附录A"):
            appendix_heading = paragraph
            break
    if appendix_heading is None:
        raise ValueError("Appendix heading not found")
    remove_between(ref_heading, appendix_heading)
    current = ref_heading
    for idx, ref in enumerate(REFERENCES, start=1):
        current = insert_paragraph_after(current, f"[{idx}] {ref}")
        format_reference(current)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    args = parser.parse_args()

    document = Document(args.input)

    h11 = find_paragraph(document, "1.1 研究背景与意义")
    h12 = find_paragraph(document, "1.2 国内外研究现状")
    h13 = find_paragraph(document, "1.3 研究目标与主要内容")
    remove_between(h11, h12)
    insert_many_after(h11, BODY_11)
    h12 = find_paragraph(document, "1.2 国内外研究现状")
    h13 = find_paragraph(document, "1.3 研究目标与主要内容")
    remove_between(h12, h13)
    insert_many_after(h12, BODY_12)

    replace_matching_paragraph(document, P13_TARGET_PREFIX, P13_REPLACEMENT)
    replace_matching_paragraph(document, P15_TARGET_PREFIX, P15_REPLACEMENT)
    replace_references(document)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    print(str(output))


if __name__ == "__main__":
    main()
