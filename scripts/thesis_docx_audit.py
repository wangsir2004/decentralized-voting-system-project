from __future__ import annotations

import argparse
import json
import re
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


CHINESE_NUM = {
    "一": 1,
    "二": 2,
    "三": 3,
    "四": 4,
    "五": 5,
    "六": 6,
    "七": 7,
    "八": 8,
    "九": 9,
    "十": 10,
}


def chinese_to_int(text: str) -> int | None:
    text = text.strip()
    if not text:
        return None
    if text in CHINESE_NUM:
        return CHINESE_NUM[text]
    if text == "十":
        return 10
    if text.startswith("十") and len(text) == 2:
        return 10 + CHINESE_NUM.get(text[1], 0)
    if text.endswith("十") and len(text) == 2:
        return CHINESE_NUM.get(text[0], 0) * 10
    if "十" in text and len(text) == 3:
        return CHINESE_NUM.get(text[0], 0) * 10 + CHINESE_NUM.get(text[2], 0)
    return None


def paragraph_has_drawing(paragraph) -> int:
    return len(paragraph._element.xpath(".//w:drawing | .//w:pict"))


def element_text(block) -> str:
    if getattr(block, "text", None) is not None:
        return " ".join(block.text.split())
    return ""


def iter_blocks(document: Document):
    paragraphs = iter(document.paragraphs)
    tables = iter(document.tables)
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", next(paragraphs)
        elif child.tag == qn("w:tbl"):
            yield "table", next(tables)


def normalize_caption(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def caption_kind(text: str) -> str | None:
    stripped = normalize_caption(text)
    if re.match(r"^图\s*(?:\d+(?:[.\-]\d+)?|[A-Z]\.?\d+)\s*[\u4e00-\u9fffA-Za-z0-9（(]", stripped):
        return "figure"
    if re.match(r"^表\s*(?:\d+(?:[.\-]\d+)?|[A-Z]\.?\d+)\s*[\u4e00-\u9fffA-Za-z0-9（(]", stripped):
        return "table"
    if re.match(r"^Figure\s+(?:\d+(?:[.\-]\d+)?|[A-Z]\.?\d+)", stripped, re.I):
        return "figure_en"
    if re.match(r"^Table\s+(?:\d+(?:[.\-]\d+)?|[A-Z]\.?\d+)", stripped, re.I):
        return "table_en"
    return None


def caption_number(text: str) -> str | None:
    stripped = normalize_caption(text)
    match = re.match(r"^(?:图|表|Figure|Table)\s*((?:\d+(?:[.\-]\d+)?)|(?:[A-Z]\.?\d+))", stripped, re.I)
    return match.group(1).replace("-", ".") if match else None


def chapter_from_text(text: str) -> tuple[int | None, str | None]:
    stripped = normalize_caption(text)
    match = re.match(r"^第([一二三四五六七八九十]+)章\s*(.+)?$", stripped)
    if match:
        return chinese_to_int(match.group(1)), stripped
    match = re.match(r"^第(\d+)章\s*(.+)?$", stripped)
    if match:
        return int(match.group(1)), stripped
    return None, None


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("--json-out")
    parser.add_argument("--text-out")
    args = parser.parse_args()

    docx_path = Path(args.docx)
    document = Document(str(docx_path))

    blocks = []
    current_chapter = None
    chapter_title = None
    paragraph_counter = 0
    table_counter = 0
    style_counter = Counter()
    direct_run_fonts = Counter()
    for index, (kind, block) in enumerate(iter_blocks(document)):
        if kind == "paragraph":
            paragraph_counter += 1
            text = element_text(block)
            style = block.style.name if block.style is not None else ""
            style_counter[style] += 1
            chapter_no, title = chapter_from_text(text)
            if chapter_no is not None:
                current_chapter = chapter_no
                chapter_title = title
            for run in block.runs:
                if run.font.name:
                    direct_run_fonts[run.font.name] += 1
            blocks.append(
                {
                    "index": index,
                    "kind": kind,
                    "text": text,
                    "style": style,
                    "drawings": paragraph_has_drawing(block),
                    "chapter": current_chapter,
                    "chapter_title": chapter_title,
                    "caption_kind": caption_kind(text),
                    "caption_number": caption_number(text),
                }
            )
        else:
            table_counter += 1
            rows = len(block.rows)
            cols = max((len(row.cells) for row in block.rows), default=0)
            sample = " | ".join(cell.text.replace("\n", " ").strip() for cell in block.rows[0].cells[:4]) if rows else ""
            blocks.append(
                {
                    "index": index,
                    "kind": kind,
                    "table_no": table_counter,
                    "rows": rows,
                    "cols": cols,
                    "sample": " ".join(sample.split())[:160],
                    "chapter": current_chapter,
                    "chapter_title": chapter_title,
                }
            )

    captions = [b for b in blocks if b.get("caption_kind")]
    figures = [b for b in blocks if b.get("kind") == "paragraph" and b.get("drawings")]
    tables = [b for b in blocks if b.get("kind") == "table"]

    by_index = {b["index"]: b for b in blocks}

    def nearest_nonempty(start: int, step: int, limit: int = 4):
        seen = 0
        i = start + step
        while i in by_index and seen < limit:
            block = by_index[i]
            if block["kind"] == "table" or block.get("text"):
                return block
            i += step
            seen += 1
        return None

    issues = []
    for table in tables:
        before = nearest_nonempty(table["index"], -1)
        after = nearest_nonempty(table["index"], 1)
        before_before = nearest_nonempty(before["index"], -1) if before else None
        has_table_caption = bool(
            before
            and (
                before.get("caption_kind") == "table"
                or (
                    before.get("caption_kind") == "table_en"
                    and before_before
                    and before_before.get("caption_kind") == "table"
                )
            )
        )
        if not has_table_caption:
            issues.append(
                {
                    "type": "table_caption_position",
                    "block": table["index"],
                    "table_no": table["table_no"],
                    "message": "表格前未发现中文表题。",
                    "context_before": before.get("text", before.get("sample", "")) if before else "",
                    "context_after": after.get("text", after.get("sample", "")) if after else "",
                }
            )
    for figure in figures:
        if figure.get("chapter") is None:
            continue
        after = nearest_nonempty(figure["index"], 1)
        before = nearest_nonempty(figure["index"], -1)
        if not after or after.get("caption_kind") != "figure":
            issues.append(
                {
                    "type": "figure_caption_position",
                    "block": figure["index"],
                    "message": "图片后未发现中文图题。",
                    "context_before": before.get("text", before.get("sample", "")) if before else "",
                    "context_after": after.get("text", after.get("sample", "")) if after else "",
                }
            )

    seq_by_chapter = defaultdict(lambda: {"figure": 0, "table": 0})
    for caption in captions:
        kind = caption["caption_kind"]
        if kind not in {"figure", "table"}:
            continue
        chapter = caption.get("chapter")
        number = caption.get("caption_number")
        if chapter is None or number is None:
            continue
        if re.match(r"^[A-Z]", number, re.I):
            continue
        seq_by_chapter[chapter][kind] += 1
        expected = f"{chapter}.{seq_by_chapter[chapter][kind]}"
        if number != expected:
            issues.append(
                {
                    "type": "caption_number",
                    "block": caption["index"],
                    "kind": kind,
                    "actual": number,
                    "expected": expected,
                    "text": caption["text"],
                }
            )

    report = {
        "docx": str(docx_path),
        "paragraphs": paragraph_counter,
        "tables": table_counter,
        "figures": len(figures),
        "captions": {
            "figure": sum(1 for c in captions if c["caption_kind"] == "figure"),
            "figure_en": sum(1 for c in captions if c["caption_kind"] == "figure_en"),
            "table": sum(1 for c in captions if c["caption_kind"] == "table"),
            "table_en": sum(1 for c in captions if c["caption_kind"] == "table_en"),
        },
        "chapters": [
            {"chapter": b["chapter"], "title": b["chapter_title"], "block": b["index"]}
            for b in blocks
            if b.get("chapter_title") and chapter_from_text(b.get("text", ""))[0] is not None
        ],
        "styles_top": style_counter.most_common(20),
        "direct_run_fonts_top": direct_run_fonts.most_common(20),
        "issues": issues,
        "caption_samples": [
            {"block": b["index"], "kind": b["caption_kind"], "chapter": b.get("chapter"), "text": b["text"]}
            for b in captions[:80]
        ],
        "table_samples": [
            {
                "block": b["index"],
                "table_no": b["table_no"],
                "chapter": b.get("chapter"),
                "rows": b["rows"],
                "cols": b["cols"],
                "sample": b["sample"],
            }
            for b in tables[:80]
        ],
    }

    if args.json_out:
        Path(args.json_out).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        f"DOCX: {docx_path}",
        f"Paragraphs: {paragraph_counter}",
        f"Tables: {table_counter}",
        f"Figures: {len(figures)}",
        f"Captions: {report['captions']}",
        "",
        "Chapters:",
    ]
    for chapter in report["chapters"]:
        lines.append(f"- block {chapter['block']}: {chapter['title']}")
    lines.extend(["", "Issues:"])
    for issue in issues[:200]:
        lines.append(f"- {issue}")
    if args.text_out:
        Path(args.text_out).write_text("\n".join(lines), encoding="utf-8")
    else:
        print("\n".join(lines))


if __name__ == "__main__":
    main()
