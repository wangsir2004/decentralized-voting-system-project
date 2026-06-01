from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


def get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_run_font(run, east="宋体", west="Times New Roman", size=Pt(12), bold=None):
    run.font.name = west
    run.font.size = size
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:cs"), west)


def clear_runs(paragraph: Paragraph):
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str):
    clear_runs(paragraph)
    if text:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.xpath(".//w:drawing | .//w:pict"))


def paragraph_has_math(paragraph: Paragraph) -> bool:
    xml = paragraph._element.xml
    return "<m:oMath" in xml or "<m:oMathPara" in xml


def remove_paragraph_numbering(paragraph: Paragraph) -> bool:
    p_pr = paragraph._p.pPr
    if p_pr is None:
        return False
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        return False
    p_pr.remove(num_pr)
    return True


def normalize_spaces(text: str) -> str:
    text = text.replace("\u3000", " ")
    text = re.sub(r"[ ]{2,}", " ", text)
    return text.strip()


def is_chapter_heading(text: str) -> bool:
    return bool(re.match(r"^第[一二三四五六七八九十\d]+章\s*.+", normalize_spaces(text)))


def is_section_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\s+.+", normalize_spaces(text)))


def is_subsection_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\d+\s+.+", normalize_spaces(text)))


def is_third_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+\.\d+\.\d+\s+.+", normalize_spaces(text)))


def is_final_heading(text: str) -> bool:
    s = normalize_spaces(text)
    return s in {"参考文献", "致谢"} or bool(re.match(r"^附录[A-Z]\s+.+", s))


def is_caption(text: str) -> str | None:
    s = normalize_spaces(text)
    if re.match(r"^图\s*\d+(?:\.\d+)?\s+.+", s):
        return "figure_cn"
    if re.match(r"^表\s*(?:\d+(?:\.\d+)?|[A-Z]\.?\d+)\s+.+", s):
        return "table_cn"
    if re.match(r"^Table\s+(?:\d+(?:\.\d+)?|[A-Z]\.?\d+)\s+.+", s, re.I):
        return "table_en"
    if re.match(r"^Figure\s+(?:\d+(?:\.\d+)?|[A-Z]\.?\d+)\s+.+", s, re.I):
        return "figure_en"
    return None


def normalize_caption(text: str) -> str:
    s = normalize_spaces(text)
    s = re.sub(r"^(图|表)\s*(\d+\.\d+)\s*", r"\1 \2 ", s)
    s = re.sub(r"^(图|表)\s*([A-Z])\.?(\d+)\s*", r"\1 \2\3 ", s)
    s = re.sub(r"^(Figure|Table)\s+([A-Z])\.?(\d+)\s*", r"\1 \2\3 ", s, flags=re.I)
    s = re.sub(r"^(Figure|Table)\s+(\d+\.\d+)\s*", r"\1 \2 ", s, flags=re.I)
    s = s.replace("图 6.4 Sepolia 部署证据表", "图 6.4 Sepolia 部署证据截图")
    return normalize_spaces(s)


def format_paragraph_runs(paragraph, east="宋体", west="Times New Roman", size=Pt(12), bold=None):
    for run in paragraph.runs:
        set_run_font(run, east=east, west=west, size=size, bold=bold)


def format_body_paragraph(paragraph):
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Pt(24)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(12), None)


def format_code_paragraph(paragraph):
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.0
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, "Consolas", "Consolas", Pt(9.5), False)


def format_caption_paragraph(paragraph, kind: str):
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.25
    if kind == "table_cn":
        pf.space_before = Pt(6)
        pf.space_after = Pt(0)
        format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(10.5), False)
    elif kind == "figure_cn":
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(10.5), False)
    else:
        pf.space_before = Pt(0)
        pf.space_after = Pt(0 if "table" in kind else 6)
        format_paragraph_runs(paragraph, "Times New Roman", "Times New Roman", Pt(10.5), False)


def format_heading(paragraph, level: str):
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    if level == "chapter":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.line_spacing = 1.25
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        format_paragraph_runs(paragraph, "黑体", "Times New Roman", Pt(18), False)
    elif level == "section":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.25
        pf.space_before = Pt(12)
        pf.space_after = Pt(12)
        format_paragraph_runs(paragraph, "黑体", "Times New Roman", Pt(16), False)
    elif level == "subsection":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.25
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(14), True)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.line_spacing = 1.25
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(12), False)


def format_toc_paragraph(paragraph):
    s = paragraph.text
    s = re.sub(r"^\uf0b7\s*\t?", "", s)
    s = re.sub(r"\t-\s*(\d+)\s*-\s*$", r"\t\1", s)
    s = re.sub(r" {2,}", " ", s)
    set_paragraph_text(paragraph, s)
    pf = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Pt(0)
    pf.line_spacing = 1.25
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if paragraph.style.name.lower().startswith("toc 2") or re.match(r"^\d+\.\d+", s):
        pf.left_indent = Pt(24)
        format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(12), False)
    else:
        pf.left_indent = Pt(0)
        format_paragraph_runs(paragraph, "黑体", "Times New Roman", Pt(12), False)
    tabs = pf.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = get_or_add(tc_pr, "w:tcMar")
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = get_or_add(tc_mar, f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_border(element, edge, value="single", size="8", color="000000"):
    borders = get_or_add(element, "w:tcBorders" if element.tag == qn("w:tcPr") else "w:tblBorders")
    tag = f"w:{edge}"
    border = borders.find(qn(tag))
    if border is None:
        border = OxmlElement(tag)
        borders.append(border)
    border.set(qn("w:val"), value)
    if value != "nil":
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def set_table_layout_fixed(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_layout_fixed(table)

    tbl_pr = table._tbl.tblPr
    borders = get_or_add(tbl_pr, "w:tblBorders")
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")

    row_count = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = tr_pr.find(qn("w:tblHeader"))
            if tbl_header is None:
                tbl_header = OxmlElement("w:tblHeader")
                tr_pr.append(tbl_header)
            tbl_header.set(qn("w:val"), "true")
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            for edge in ["top", "left", "bottom", "right"]:
                set_border(tc_pr, edge, "nil")
            if r_idx == 0:
                set_border(tc_pr, "top", "single", "12")
                set_border(tc_pr, "bottom", "single", "8")
            if r_idx == row_count - 1:
                set_border(tc_pr, "bottom", "single", "12")
            for paragraph in cell.paragraphs:
                pf = paragraph.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.line_spacing = 1.25
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                if r_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(10.5), True)
                else:
                    if c_idx == 0 or len(paragraph.text) > 18:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(10.5), False)


def is_code_like(text: str) -> bool:
    s = text.strip()
    return (
        "\n" in text
        or s.startswith("require(")
        or s.startswith("npm ")
        or s.startswith("slither ")
        or s.startswith("bytes32 ")
        or s.startswith("voted[")
    )


def format_front_matter(paragraphs, report):
    replacements = {
        "原创性声明（宋体三号）": "原创性声明",
        "关于毕业设计说明书版权使用授权的说明（宋体三号）": "关于毕业设计说明书版权使用授权的说明",
        "Key words：": "Key words: ",
        "基于Solidity智能合约与Ethereum测试网": "基于 Solidity 智能合约与 Ethereum 测试网",
        "Solidity合约": "Solidity 合约",
    }
    for paragraph in paragraphs:
        text = paragraph.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text.startswith("Key words:"):
            new_text = re.sub(r"^Key words:\s*", "Key words: ", new_text)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)
            report["text_replacements"].append({"old": text, "new": new_text})

        s = normalize_spaces(paragraph.text)
        if s in {"原创性声明", "关于毕业设计说明书版权使用授权的说明"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            format_paragraph_runs(paragraph, "宋体", "Times New Roman", Pt(16), False)
        elif s == "摘 要":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            format_paragraph_runs(paragraph, "黑体", "Times New Roman", Pt(18), False)
        elif s == "ABSTRACT":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(12)
            format_paragraph_runs(paragraph, "Times New Roman", "Times New Roman", Pt(18), False)


def add_toc_title_if_missing(document, report):
    if any(normalize_spaces(p.text).replace(" ", "") == "目录" for p in document.paragraphs):
        return
    first_toc = None
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.style.name.lower().startswith("toc") or paragraph.text.startswith("\uf0b7"):
            first_toc = i
            break
    if first_toc is None:
        return

    previous = document.paragraphs[first_toc - 1]
    title_para = insert_paragraph_after(previous, "目  录")
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.first_line_indent = Pt(0)
    title_para.paragraph_format.line_spacing = 1.25
    title_para.paragraph_format.space_before = Pt(12)
    title_para.paragraph_format.space_after = Pt(12)
    format_paragraph_runs(title_para, "黑体", "Times New Roman", Pt(18), False)
    report["toc_title_inserted"] = True


def optimize_docx(source: Path, output: Path) -> dict:
    document = Document(str(source))
    report = {
        "source": str(source),
        "output": str(output),
        "text_replacements": [],
        "captions_normalized": [],
        "toc_title_inserted": False,
        "toc_entries_normalized": 0,
        "numbering_removed": 0,
        "tables_formatted": 0,
        "images_resized": 0,
    }

    add_toc_title_if_missing(document, report)
    format_front_matter(document.paragraphs, report)

    first_chapter_seen = False
    for paragraph in document.paragraphs:
        if remove_paragraph_numbering(paragraph):
            report["numbering_removed"] += 1
        text = paragraph.text
        stripped = normalize_spaces(text)

        if paragraph.style.name.lower().startswith("toc") or paragraph.text.startswith("\uf0b7"):
            format_toc_paragraph(paragraph)
            report["toc_entries_normalized"] += 1
            continue

        caption_kind = is_caption(text)
        if caption_kind:
            normalized = normalize_caption(text)
            if normalized != text:
                set_paragraph_text(paragraph, normalized)
                report["captions_normalized"].append({"old": text, "new": normalized})
            format_caption_paragraph(paragraph, caption_kind)
            continue

        if paragraph_has_drawing(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.25
            continue

        if paragraph_has_math(paragraph):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            paragraph.paragraph_format.line_spacing = 1.25
            format_paragraph_runs(paragraph, "Cambria Math", "Cambria Math", Pt(12), False)
            continue

        if is_chapter_heading(stripped) or is_final_heading(stripped):
            first_chapter_seen = first_chapter_seen or is_chapter_heading(stripped)
            if stripped != text:
                set_paragraph_text(paragraph, stripped)
            format_heading(paragraph, "chapter")
            continue

        if first_chapter_seen:
            if is_third_heading(stripped):
                if stripped != text:
                    set_paragraph_text(paragraph, stripped)
                format_heading(paragraph, "third")
            elif is_subsection_heading(stripped):
                if stripped != text:
                    set_paragraph_text(paragraph, stripped)
                format_heading(paragraph, "subsection")
            elif is_section_heading(stripped) or re.match(r"^[A-Z]\.\d+\s+.+", stripped):
                if stripped != text:
                    set_paragraph_text(paragraph, stripped)
                format_heading(paragraph, "section")
            elif is_code_like(text):
                format_code_paragraph(paragraph)
            elif stripped:
                format_body_paragraph(paragraph)

    for shape_index, shape in enumerate(document.inline_shapes, start=1):
        if shape_index == 1:
            continue
        max_width = Cm(14.6)
        max_height = Cm(14.2)
        scale = min(max_width / shape.width, max_height / shape.height, 1)
        if scale < 1:
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)
            report["images_resized"] += 1

    for table in document.tables:
        format_table(table)
        report["tables_formatted"] += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))
    return report


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("output")
    parser.add_argument("--report")
    args = parser.parse_args()

    report = optimize_docx(Path(args.source), Path(args.output))
    if args.report:
        Path(args.report).write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
