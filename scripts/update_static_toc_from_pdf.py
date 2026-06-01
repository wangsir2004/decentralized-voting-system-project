from __future__ import annotations

import argparse
import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def clear_paragraph_keep_properties(paragraph):
    for child in list(paragraph._element):
        if child.tag != qn("w:pPr"):
            paragraph._element.remove(child)


def set_run_font(run, east="宋体", west="Times New Roman", size=Pt(12), bold=None):
    run.font.name = west
    run.font.size = size
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        from docx.oxml import OxmlElement

        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:cs"), west)


def normalize_text(text: str) -> str:
    text = re.sub(r"\s+", "", text)
    text = text.replace("，", ",")
    return text


def extract_pdf_pages(pdf: Path, pdftotext: str) -> list[str]:
    info = subprocess.run(
        ["pdfinfo", str(pdf)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="ignore",
        check=True,
    )
    match = re.search(r"^Pages:\s+(\d+)", info.stdout, re.M)
    if not match:
        raise RuntimeError("Could not determine PDF page count")
    page_count = int(match.group(1))
    pages: list[str] = []
    for page in range(1, page_count + 1):
        result = subprocess.run(
            [pdftotext, "-f", str(page), "-l", str(page), "-layout", str(pdf), "-"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
            check=True,
        )
        pages.append(result.stdout)
    return pages


def visible_footer_number(page_text: str, fallback: int) -> int:
    for line in reversed([line.strip() for line in page_text.splitlines()]):
        if re.fullmatch(r"\d{1,3}", line):
            return int(line)
    return fallback


def parse_toc_entry(text: str) -> tuple[str, str] | None:
    text = text.replace("\uf0b7", "").strip()
    if not text:
        return None
    if "\t" in text:
        title, old_page = text.rsplit("\t", 1)
        title = title.strip()
        old_page = old_page.strip()
        if title and old_page.isdigit():
            return title, old_page
    match = re.match(r"^(.*?)\s+(\d+)$", text)
    if match:
        return match.group(1).strip(), match.group(2)
    return None


def update_toc(docx: Path, pdf: Path, out: Path, min_pdf_page: int, pdftotext: str) -> dict:
    pages = extract_pdf_pages(pdf, pdftotext)
    return update_toc_from_pages(docx, pages, out, min_pdf_page)


def update_toc_from_pages(docx: Path, pages: list[str], out: Path, min_pdf_page: int) -> dict:
    page_norm = [normalize_text(page) for page in pages]
    footers = [visible_footer_number(page, i + 1) for i, page in enumerate(pages)]

    document = Document(str(docx))
    updates = []
    for paragraph in document.paragraphs:
        if not paragraph.style.name.lower().startswith("toc"):
            continue
        parsed = parse_toc_entry(paragraph.text)
        if not parsed:
            continue
        title, old_page = parsed
        needle = normalize_text(title)
        actual_page = None
        for pdf_index in range(max(min_pdf_page - 1, 0), len(page_norm)):
            if needle in page_norm[pdf_index]:
                actual_page = footers[pdf_index]
                break
        if actual_page is None:
            actual_page = int(old_page)
        new_text = f"{title}\t{actual_page}"
        clear_paragraph_keep_properties(paragraph)
        run = paragraph.add_run(new_text)
        if paragraph.style.name.lower().startswith("toc 2") or re.match(r"^\d+\.\d+|^[A-Z]\.\d+", title):
            paragraph.paragraph_format.left_indent = Pt(24)
            set_run_font(run, "宋体", "Times New Roman", Pt(12), False)
        else:
            paragraph.paragraph_format.left_indent = Pt(0)
            set_run_font(run, "黑体", "Times New Roman", Pt(12), False)
        paragraph.paragraph_format.tab_stops.clear_all()
        paragraph.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), WD_ALIGN_PARAGRAPH.RIGHT)
        updates.append({"title": title, "old": int(old_page), "new": int(actual_page)})

    document.save(str(out))
    return {"updated": len(updates), "updates": updates}


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx")
    parser.add_argument("pdf")
    parser.add_argument("out")
    parser.add_argument("--min-pdf-page", type=int, default=8)
    parser.add_argument("--pdftotext", default="pdftotext")
    parser.add_argument("--pdf-text")
    parser.add_argument("--report")
    args = parser.parse_args()
    if args.pdf_text:
        pages = Path(args.pdf_text).read_text(encoding="utf-8", errors="ignore").split("\f")
        result = update_toc_from_pages(Path(args.docx), pages, Path(args.out), args.min_pdf_page)
    else:
        result = update_toc(Path(args.docx), Path(args.pdf), Path(args.out), args.min_pdf_page, args.pdftotext)
    if args.report:
        Path(args.report).write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"updated {result['updated']} toc entries")


if __name__ == "__main__":
    main()
