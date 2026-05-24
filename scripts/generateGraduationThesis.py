from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
TEMPLATE = DESKTOP / "毕业设计" / "4.毕业设计说明书格式新的模板.docx"
OUT_DOCX = ROOT / "deliverables" / "graduation" / "docx" / "王靖宇-毕业设计说明书-去中心化电子投票系统.docx"
FIG = ROOT / "deliverables" / "graduation" / "figures"
SHOT = ROOT / "deliverables" / "graduation" / "screenshots"

TITLE = "基于 Solidity 智能合约与 Ethereum 测试网的去中心化电子投票系统设计与实现"
AUTHOR = "王靖宇"
STUDENT_ID = "225288"
COLLEGE = "人工智能与数据科学学院"
MAJOR = "计算机科学与技术"
CLASS_NAME = "计算机224"
ADVISOR = "王立鹏"
ADVISOR_TITLE = "讲师"
REVIEWER = "闫文杰"
REVIEWER_TITLE = "副教授"
HEADER_TEXT = "河北工业大学2026届本科毕业设计说明书"


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: int | None = 10.5, bold: bool = False):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line: bool = True):
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(end)


def add_paragraph_bottom_border(paragraph, color: str = "9AA6B2") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def clear_story(story) -> None:
    for paragraph in story.paragraphs:
        paragraph.text = ""


def add_header_text(header) -> None:
    clear_story(header)
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(HEADER_TEXT)
    set_run_font(run, east_asia="宋体", size=9)
    add_paragraph_bottom_border(p)


def add_footer_page_number(footer) -> None:
    clear_story(footer)
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    set_run_font(run, size=9)
    add_field(run, "PAGE")


def configure_section(
    section,
    *,
    top: float,
    bottom: float,
    left: float,
    right: float,
    header_distance: float,
    footer_distance: float,
    different_first: bool = False,
    page_number: bool = False,
) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)
    section.header_distance = Cm(header_distance)
    section.footer_distance = Cm(footer_distance)
    section.different_first_page_header_footer = different_first

    for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
        story.is_linked_to_previous = False
        clear_story(story)

    add_header_text(section.header)
    if different_first:
        clear_story(section.first_page_header)

    if page_number:
        add_footer_page_number(section.footer)


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    configure_section(
        section,
        top=2.54,
        bottom=2.54,
        left=2.5,
        right=1.8,
        header_distance=1.5,
        footer_distance=1.75,
        different_first=True,
        page_number=False,
    )

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.font.bold = None

    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 3"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def start_declaration_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(
        section,
        top=3.4,
        bottom=2.5,
        left=2.7,
        right=2.7,
        header_distance=1.5,
        footer_distance=1.75,
        page_number=False,
    )


def start_main_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(
        section,
        top=3.4,
        bottom=2.5,
        left=2.7,
        right=2.7,
        header_distance=2.5,
        footer_distance=2.2,
        page_number=True,
    )


def add_center(doc: Document, text: str, size: int = 10.5, bold: bool = False, font_name: str = "宋体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, east_asia=font_name, size=size, bold=bold)
    return p


def add_para(doc: Document, text: str, *, first_line: bool = True, bold_prefix: str | None = None):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, bold=True)
        run2 = p.add_run(text[len(bold_prefix):])
        set_run_font(run2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_mono(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for line in text.splitlines():
        run = p.add_run(line)
        set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=9.5)
        run.add_break()
    p_pr = p._p.get_or_add_pPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F5F7FA")
    p_pr.append(shade)
    return p


def add_chapter(doc: Document, text: str, *, break_before: bool = True):
    if break_before and len(doc.paragraphs) > 0:
        doc.add_page_break()
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, east_asia="黑体", ascii_font="Arial", size=15, bold=True)
    return p


def add_section_heading(doc: Document, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12 if level == 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    for run in p.runs:
        set_run_font(run, east_asia="黑体", ascii_font="Arial", size=14 if level == 2 else 12, bold=True)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=9)
    return p


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell(cell, text: str, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "9AA6B2")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    set_table_borders(table)
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "DDE7EF")
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shade)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cells[i], value, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_picture(doc: Document, image_path: Path, caption: str, width: float = 5.8):
    if not image_path.exists():
        add_para(doc, f"（图片缺失：{image_path.name}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, caption)


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "毕业设计说明书", size=38, bold=False, font_name="方正大标宋简体")
    doc.add_paragraph()
    add_center(doc, "题       目：   基于 Solidity 智能合约与 Ethereum 测试网", size=14, font_name="楷体")
    add_center(doc, "去中心化电子投票系统设计与实现", size=14, font_name="楷体")
    add_center(doc, f"学       院：{COLLEGE}", size=14, font_name="楷体")
    add_center(doc, f"专       业：{MAJOR}", size=14, font_name="楷体")
    add_center(doc, f"作       者：姓名 {AUTHOR}  学号 {STUDENT_ID}", size=14, font_name="楷体")
    add_center(doc, f"指 导 教 师：姓名 {ADVISOR}  职称 {ADVISOR_TITLE}", size=14, font_name="楷体")
    add_center(doc, f"评   阅  者：姓名 {REVIEWER}  职称 {REVIEWER_TITLE}", size=14, font_name="楷体")
    for _ in range(7):
        doc.add_paragraph()
    add_center(doc, "二〇二六 年 六 月", size=12, font_name="楷体")


def add_declarations(doc: Document):
    add_center(doc, "原创性声明", size=16, bold=True, font_name="宋体")
    add_para(doc, "本人郑重声明：所呈交的毕业设计说明书，是本人在导师指导下，进行研究工作所取得的成果。除文中已经注明引用的内容外，本设计不包含任何他人或集体已经发表的作品内容，也不包含本人为获得其他学位而使用过的材料。对本设计所涉及的研究工作做出贡献的其他个人或集体，均已在文中以明确方式标明。本毕业设计说明书原创性声明的法律责任由本人承担。")
    doc.add_paragraph()
    add_para(doc, "作者签名：                      日期：      年    月    日", first_line=False)
    doc.add_paragraph()
    add_center(doc, "关于毕业设计说明书版权使用授权的说明", size=16, bold=True, font_name="宋体")
    add_para(doc, "本人完全了解河北工业大学关于收集、保存、使用毕业设计说明书的以下规定：本科生在校攻读学位期间毕业设计工作的知识产权单位属河北工业大学，学校有权采用影印、缩印、扫描、数字化或其他手段保存毕业设计说明书；学校有权提供本设计全文或者部分内容的阅览服务；学校有权将毕业设计说明书的全部或部分内容编入有关数据库进行检索、交流；学校有权向国家有关部门或者机构送交毕业设计说明书的复印件和电子版。")
    add_para(doc, "（保密的毕业设计说明书在解密后适用本授权说明）")
    doc.add_paragraph()
    add_para(doc, "作  者  签  名：                 日期：      年    月    日", first_line=False)
    add_para(doc, "导  师  签  名：                 日期：      年    月    日", first_line=False)


def add_abstracts(doc: Document):
    add_center(doc, "摘　　要", size=18, bold=True, font_name="黑体")
    add_para(doc, "随着在线治理、数字化协同和远程办公场景不断扩展，电子投票系统需要在可用性、透明性、结果可复核和安全性之间取得平衡。传统中心化电子投票系统通常由单一服务端保存选民状态和计票结果，容易面临单点故障、内部人为干预、过程追溯困难和结果公信力不足等问题。针对上述问题，本文设计并实现了一种基于 Solidity 智能合约与 Ethereum Sepolia 测试网的去中心化电子投票系统。系统以 VotingSystem 智能合约承载投票主题、候选项、投票时间、白名单资格校验、防重复投票和链上计票等核心逻辑；采用 Merkle Tree 保存白名单根以降低链上存储成本；使用 React、Vite、TypeScript、Ethers.js 和 MetaMask 构建前端 DApp，实现钱包连接、资格证明读取、投票提交、交易回执展示和实时结果查询。")
    add_para(doc, "本文完成了需求分析、系统设计、合约实现、前端开发、自动化测试、Gas 消耗分析、安全审计和 Sepolia 测试网部署。测试结果显示，系统自动化测试共 24 个用例全部通过；本地部署 Gas 为 767,754，vote 方法平均 Gas 为 73,661；最新 Sepolia 部署合约地址为 0x506db4C4b9A63d127b3613AC87B2Cae3Fe185d2b，投票窗口约为 3 天。Slither 审计结果仅提示与投票截止时间相关的 timestamp 风险，该风险属于可解释的业务边界。实践结果表明，系统能够实现地址级资格控制、链上公开计票和结果可追溯，为小规模电子投票场景提供了一种可运行、可验证的实现方案。")
    add_para(doc, "关键词：去中心化投票；智能合约；Solidity；Merkle Tree；Ethereum；DApp", first_line=False, bold_prefix="关键词：")
    doc.add_page_break()

    add_center(doc, "ABSTRACT", size=18, bold=True, font_name="Times New Roman")
    add_para(doc, "With the expansion of online governance, digital collaboration and remote working scenarios, electronic voting systems need to balance usability, transparency, verifiability and security. Traditional centralized electronic voting systems usually store voter states and tallying results on a single server, which may lead to single point of failure, internal manipulation, weak process traceability and insufficient public trust. To address these problems, this thesis designs and implements a decentralized electronic voting system based on Solidity smart contracts and the Ethereum Sepolia test network. The VotingSystem smart contract is used to manage the voting title, candidates, voting period, whitelist verification, duplicate-vote prevention and on-chain tallying. A Merkle Tree whitelist is adopted so that only the Merkle root is stored on chain, reducing storage overhead. The front-end DApp is implemented with React, Vite, TypeScript, Ethers.js and MetaMask, supporting wallet connection, proof loading, vote submission, transaction receipt display and real-time result query.")
    add_para(doc, "The work covers requirement analysis, system design, smart contract implementation, front-end development, automated testing, gas consumption analysis, security auditing and deployment on Sepolia. The automated test suite contains 24 passed cases. The local deployment consumes 767,754 gas, while the average gas consumption of the vote method is 73,661. The latest deployed Sepolia contract address is 0x506db4C4b9A63d127b3613AC87B2Cae3Fe185d2b, with a voting window of about three days. Slither reports only timestamp-related warnings for the voting deadline, which are treated as acceptable business constraints. The results show that the system can provide address-level eligibility control, transparent on-chain tallying and traceable voting evidence, offering an executable and verifiable prototype for small-scale electronic voting scenarios.")
    add_para(doc, "Key words: decentralized voting; smart contract; Solidity; Merkle Tree; Ethereum; DApp", first_line=False, bold_prefix="Key words:")
    doc.add_page_break()


def add_toc(doc: Document):
    add_center(doc, "目　　录", size=18, bold=True, font_name="黑体")
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    add_field(run, r'TOC \o "1-3" \h \z \u')


def add_body(doc: Document):
    add_chapter(doc, "第一章  绪论", break_before=False)
    add_section_heading(doc, "1.1  研究背景与意义")
    add_para(doc, "电子投票是信息化治理、组织决策和在线协作中的重要基础能力。与纸质投票相比，电子投票能够降低人工计票成本，提高结果统计效率，并支持远程参与。然而，传统电子投票系统通常以中心化服务器作为唯一可信执行主体，选民名单、投票状态和计票结果均由服务端数据库保存。当系统管理员权限过大、服务器遭受攻击或日志记录不完整时，投票过程可能难以被外部独立验证。")
    add_para(doc, "区块链技术通过分布式账本、密码学签名和不可篡改的交易记录，为投票系统提供了新的实现思路。将投票规则写入智能合约后，候选项、资格校验、重复投票限制和计票逻辑可以在链上按照公开代码自动执行。任何观察者都可以根据合约地址、交易 Hash 和事件日志复核投票过程，从而降低对中心化平台的单点信任依赖。")
    add_para(doc, "本文课题围绕“基于 Solidity 智能合约与 Ethereum 测试网的去中心化电子投票系统设计与实现”展开，目标不是构建真实政务选举系统，而是完成一个面向毕业设计和教学演示的小规模可执行原型。该原型重点验证智能合约投票、Merkle Tree 白名单、钱包签名、Sepolia 部署、链上结果查询和安全审计等关键技术路径。")
    add_section_heading(doc, "1.2  国内外研究现状")
    add_para(doc, "电子投票研究长期关注身份认证、投票匿名性、可验证计票、抗篡改和系统可用性等问题。传统方案多依赖中心化身份认证、数据库事务和审计日志来保障投票过程完整性。这类方案在工程实现上较成熟，但其可信基础仍集中在系统运营方和数据库管理方，外部审查者通常只能依赖事后报告或抽样日志。")
    add_para(doc, "近年来，区块链和智能合约被广泛用于可验证流程的原型研究。Ethereum 提供账户模型、交易签名、事件日志和合约虚拟机，使投票规则可以被编译为链上程序执行。基于智能合约的投票系统通常具有结果公开、流程可追溯、规则可检查等优势，但也存在交易成本、隐私保护、链上吞吐量和身份绑定等限制。因此，在毕业设计场景中，应明确系统边界，将其定位为小规模投票演示与技术验证，而不是直接替代真实大规模选举。")
    add_para(doc, "从国外研究和开源实践看，区块链投票系统大多围绕“可公开验证”和“隐私保护”两个方向展开。前者强调任何人都能够检查投票合约、交易日志和最终结果，适合组织内部表决、社区治理和小规模实验；后者则需要结合零知识证明、同态加密、盲签名或承诺揭示协议，技术复杂度较高，对实现正确性和审计能力要求更强。本文选择前者作为实现重点，是因为其更适合本科毕业设计周期内完成可运行原型。")
    add_para(doc, "从国内应用背景看，电子投票常出现在高校评选、社团决策、企业内部意见征集和线上活动评审等场景。这些场景通常规模有限，但对过程公开、结果可信和操作便捷有一定要求。本文系统面向这类教学和演示场景，将核心目标放在防重复投票、防非授权地址投票、链上公开计票和部署证据留存，而不直接处理政务选举中的实名身份核验、强匿名性和法律效力问题。")
    add_section_heading(doc, "1.3  研究目标与主要内容")
    add_para(doc, "本文的研究目标是设计并实现一个可运行、可部署、可测试、可追溯的去中心化电子投票系统。系统应支持管理员配置投票主题和候选项，支持使用地址白名单限制投票资格，支持选民通过钱包签名提交投票交易，支持合约自动拒绝非白名单地址、重复投票和超时投票，并支持前端读取链上结果。")
    add_para(doc, "围绕上述目标，本文完成以下工作：第一，设计投票系统的角色、功能需求、非功能需求和威胁边界；第二，基于 Solidity 编写 VotingSystem 智能合约，并使用 OpenZeppelin MerkleProof 完成白名单验证；第三，使用 Hardhat 搭建编译、测试、部署和 Gas 统计流程；第四，使用 React、TypeScript、Ethers.js 和 MetaMask 实现前端 DApp；第五，在 Sepolia 测试网上重新部署合约并同步前端配置；第六，使用 Slither 完成合约静态分析并形成审计记录。")
    add_section_heading(doc, "1.4  论文结构安排")
    add_para(doc, "全文共分为七章。第一章介绍研究背景、研究现状、研究目标和论文结构。第二章说明 Ethereum、Solidity、ECDSA、Keccak-256、Merkle Tree 和 DApp 钱包交互等相关技术。第三章给出系统需求分析和威胁模型。第四章描述系统总体架构、业务流程、投票流程、合约接口和前端模块。第五章介绍智能合约、白名单脚本、部署脚本和前端界面的具体实现。第六章给出测试结果、Gas 消耗、安全审计和 Sepolia 部署证据。第七章总结已完成工作并分析后续改进方向。")
    add_section_heading(doc, "1.5  技术路线与可行性分析")
    add_para(doc, "本文采用“需求分析—合约建模—本地测试—测试网部署—前端联调—材料归档”的技术路线。首先根据投票场景抽象出管理员、选民和审查者三类角色，并将核心约束转化为合约状态变量和函数前置条件；其次使用 Hardhat 在本地链上反复验证投票资格、重复投票和截止时间等边界；然后将合约部署到 Sepolia 测试网，通过真实交易 Hash 和区块链浏览器形成可追溯证据；最后将合约地址、ABI、白名单和运行截图整理为毕业设计材料。")
    add_para(doc, "从技术可行性看，Ethereum 账户模型能够天然区分不同投票地址，智能合约能够在链上保存不可抵赖的状态更新记录，Merkle Tree 能够用一个根哈希表示完整白名单，React 前端能够通过钱包插件完成用户交互。因此，课题所需的关键技术均有成熟生态支撑。项目没有设计自定义密码学算法，而是使用 Solidity 语言特性、OpenZeppelin 安全库和 Ethers.js 通用接口完成实现，这有助于降低毕业设计阶段的实现风险。")
    add_para(doc, "从工程可行性看，系统采用 monorepo 方式组织合约、脚本、测试和前端，所有关键步骤均可以通过 npm 命令复现。部署记录、Gas 结果、Slither 输出和前端截图统一保存到 deliverables 目录，使论文中的数据不依赖人工描述。该方式符合软件工程中“过程可记录、结果可验证、材料可追溯”的基本要求，也便于后续答辩时现场展示。")

    add_chapter(doc, "第二章  相关技术基础")
    add_section_heading(doc, "2.1  Ethereum 与智能合约")
    add_para(doc, "Ethereum 是支持智能合约的公有区块链平台。用户通过外部账户发起交易，交易经签名后广播到网络，由验证者打包进入区块。智能合约部署后拥有确定的合约地址，其代码和存储状态由链上节点共同维护。对于投票系统而言，合约可以作为公开的规则执行主体，保证投票限制和计票逻辑不依赖单一后端服务器。")
    add_para(doc, "本文选用 Sepolia 测试网进行部署。Sepolia 与主网具有相近的账户、交易和合约执行模型，但使用测试币支付 Gas，适合课程设计和毕业设计验证。系统最终部署合约地址、部署交易 Hash 和投票交易 Hash 均保存到部署记录中，便于复核。")
    add_section_heading(doc, "2.2  Solidity、Hardhat 与 OpenZeppelin")
    add_para(doc, "Solidity 是 Ethereum 生态中最常用的智能合约开发语言。本文合约使用 Solidity 0.8.24 编写，该版本内置整数溢出检查，减少了早期合约中常见的算术安全问题。Hardhat 用于合约编译、单元测试、部署脚本执行和 Gas 统计，能够将合约开发、测试和部署过程组织为可重复命令。")
    add_para(doc, "OpenZeppelin 是智能合约安全组件库。本文未自行实现 Merkle Proof 验证算法，而是使用 OpenZeppelin 的 MerkleProof 库完成 proof 校验，降低了手写密码学验证逻辑出错的概率。")
    add_section_heading(doc, "2.3  ECDSA 签名与地址机制")
    add_para(doc, "Ethereum 外部账户由私钥、公钥和地址构成。设 secp256k1 椭圆曲线的基点为 G，阶为 n，用户私钥为 d，且 1 ≤ d < n，则对应公钥为 Q = dG。Ethereum 地址不是完整公钥，而是对公钥进行 Keccak-256 哈希后取后 20 字节得到，因此地址长度为 20 字节，即 160 bit。")
    add_mono(doc, "Q = dG\naddress = last_20_bytes(keccak256(publicKey))")
    add_para(doc, "用户通过 MetaMask 发起投票交易时，钱包会对交易摘要 z 进行 ECDSA 签名，得到 (r, s, v)。节点在交易进入 EVM 执行前完成签名恢复和账户校验，合约中读取到的 msg.sender 已是验证后的交易发送者。因此，本系统在合约中不需要重复实现 ECDSA 验证，只需要基于 msg.sender 生成白名单叶子节点并校验 Merkle Proof。")
    add_mono(doc, "w  = s^(-1) mod n\nu1 = z * w mod n\nu2 = r * w mod n\nR  = u1 * G + u2 * Q\naccept <=> r == x(R) mod n")
    add_para(doc, "ECDSA 的安全性建立在椭圆曲线离散对数问题之上。已知基点 G 和公钥 Q = dG 时，攻击者若要恢复私钥 d，需要在 secp256k1 曲线上求解离散对数问题。在当前公开认知下，该问题不存在可在现实时间内完成的通用经典算法。因此，只要用户妥善保存私钥，其他账户无法伪造其交易签名，也无法冒充该地址完成投票。")
    add_para(doc, "Ethereum 交易签名并不是对明文投票选项进行简单签名，而是对经过编码和哈希后的交易数据签名。交易数据包含 nonce、gas 参数、接收地址、value、data 和 chainId 等字段。chainId 的存在可以防止同一笔签名在其他兼容链上被重放；nonce 则保证同一账户发出的交易具有严格顺序，避免旧交易被重复提交。")
    add_para(doc, "在本系统中，候选项选择和 Merkle Proof 被编码到交易 data 字段中，钱包在用户确认后对完整交易摘要签名。节点验证签名后，EVM 执行合约函数，此时 msg.sender 对应签名账户。也就是说，合约中的资格校验对象不是前端传入的任意地址，而是交易发送者本身，这一点可以避免用户通过修改前端参数冒用其他白名单地址。")
    add_para(doc, "地址长度为 20 字节而哈希输出为 32 字节，是因为二者承担的功能不同。地址用于账户标识和状态索引，160 bit 已能提供足够大的地址空间；Keccak-256 输出用于完整性校验和碰撞抵抗，256 bit 的输出空间能够显著降低不同输入映射到同一哈希值的概率。本文在白名单中使用 32 字节叶子节点，是为了使 Merkle Tree 的内部节点和根节点都保持统一的 bytes32 类型。")
    add_section_heading(doc, "2.4  Keccak-256 与 Merkle Tree")
    add_para(doc, "Keccak-256 是 Ethereum 中常用的哈希函数，输出长度为 32 字节，即 256 bit。本系统使用 Keccak-256 对白名单地址进行哈希，得到 Merkle Tree 的叶子节点。对于 k 个不同输入，哈希碰撞概率可用生日悖论近似分析。输出空间大小 N = 2^256，不发生碰撞的概率为：")
    add_mono(doc, "P(no collision) = Π(i = 0 to k - 1) (1 - i / 2^256)\nP(collision) = 1 - Π(i = 0 to k - 1) (1 - i / 2^256)\n当 k << 2^128 时，P(collision) ≈ k(k - 1) / 2^257")
    add_para(doc, "本系统白名单地址数为 k = 3，因此碰撞概率近似为 3 / 2^256，约等于 2.59 × 10^(-77)。即使扩展到 1,000,000 个地址，碰撞概率也约为 4.32 × 10^(-66)，在工程实践中可视为计算上不可行。")
    add_para(doc, "Merkle Tree 通过树形哈希结构将多个叶子节点压缩为一个根哈希。智能合约只保存 Merkle Root，前端在投票时提交当前地址对应的 proof。链上验证复杂度约为 O(log n)，相比将完整白名单写入合约并遍历验证的 O(n) 方案，能够显著降低链上存储和验证成本。")
    add_para(doc, "Merkle Tree 的验证过程可以理解为从叶子节点向根节点逐层还原。设当前地址计算出的叶子为 h0，proof 中第 i 个兄弟节点为 si，则验证器根据节点排序规则计算 hi+1 = H(sort(hi, si))。当所有 proof 节点依次计算完成后，若最终得到的根哈希等于合约中保存的 merkleRoot，则说明该叶子确实属于部署前生成的白名单集合。")
    add_para(doc, "本文脚本和合约均采用相同的叶子计算方式，即 keccak256(abi.encodePacked(address))。abi.encodePacked 会将地址按 20 字节紧凑编码，然后再进行 Keccak-256 哈希。若脚本端和合约端使用不同编码方式，例如脚本使用字符串形式地址而合约使用 address 二进制编码，则同一个地址会得到不同叶子节点，最终导致 proof 校验失败。因此，编码一致性是 Merkle 白名单能否正确运行的关键。")
    add_para(doc, "当白名单数量不是 2 的整数次幂时，Merkle Tree 生成库通常会按照既定规则处理奇数节点，例如复制末尾节点或提升节点。本项目使用脚本统一生成根和 proof，并在前端直接使用脚本输出结果，避免人工手算树结构带来的不一致。对于当前 3 个白名单地址，根哈希由 3 个叶子经过两层合并得到；第 3 个地址的 proof 深度较短，是由树生成规则和节点位置共同决定的。")
    add_para(doc, "从链上成本角度看，Merkle Proof 验证会随着 proof 长度增加而增加少量哈希计算 Gas，但不会引入链上数组存储。若白名单规模扩大到 1024 个地址，proof 长度约为 10 个 bytes32 节点；若扩大到 1,048,576 个地址，proof 长度约为 20 个节点。相比把百万级地址写入合约状态，Merkle Root 方案在部署成本和状态膨胀方面具有明显优势。")
    add_para(doc, "从安全边界看，Merkle Root 只能证明某地址在预先承诺的集合中，不能证明该地址对应真实自然人身份。因此，本文将它作为投票资格控制机制，而不是实名身份认证机制。真实场景中还需要由链下机构完成身份审核、凭证发放和地址绑定，链上合约只验证凭证或 proof 是否属于授权集合。")
    add_picture(doc, FIG / "merkle-proof-flow.png", "图2.1  Merkle Tree 白名单验证流程图", width=5.9)
    add_section_heading(doc, "2.5  DApp 前端与钱包交互")
    add_para(doc, "去中心化应用前端通常不直接保管用户私钥，而是通过浏览器钱包完成账户授权和交易签名。本文前端使用 Ethers.js 封装合约读写操作，读取类方法可通过浏览器钱包或 Sepolia 只读 RPC 获取链上状态，写入类方法必须由用户钱包签名后广播交易。该设计避免前端保存私钥，同时保留用户对交易提交的确认权。")
    add_section_heading(doc, "2.6  技术选型依据")
    add_para(doc, "合约开发环境选择 Hardhat，主要原因是其与 TypeScript、Ethers.js、Mocha/Chai 测试框架集成较好，能够同时满足合约编译、单元测试、部署脚本和 Gas 统计需求。与仅在 Remix 中手工部署相比，Hardhat 更适合保留工程过程证据，能够通过脚本重复生成部署结果和测试输出。")
    add_para(doc, "前端框架选择 React、Vite 和 TypeScript。React 组件模型适合将钱包状态、投票面板、结果统计和安全证明拆分为独立模块；Vite 启动速度较快，适合本地联调；TypeScript 能够在编译阶段发现配置字段缺失、合约返回值类型不匹配和前端状态使用错误。对于毕业设计而言，这种组合既能保证开发效率，也能体现现代 Web 工程实践。")
    add_para(doc, "白名单方案没有采用链上数组遍历，而是采用 Merkle Tree。若直接把所有选民地址写入合约数组，并在投票时遍历判断资格，则验证成本会随地址数量线性增长；若使用 mapping(address => bool) 逐个写入白名单，则部署和维护成本会随地址数量增加。Merkle Root 方案只在链上保存一个 32 字节根哈希，投票时提交 proof，由合约执行对数级验证，更适合展示区块链系统中“链下计算、链上验证”的设计思想。")
    add_table(doc, ["技术", "选型", "理由"], [
        ["合约语言", "Solidity 0.8.24", "内置溢出检查，Ethereum 生态兼容性好"],
        ["开发框架", "Hardhat", "支持测试、部署、Gas 报告和脚本化流程"],
        ["安全库", "OpenZeppelin", "复用成熟 MerkleProof 实现，减少手写风险"],
        ["前端框架", "React + Vite", "适合组件化界面与快速本地调试"],
        ["链交互库", "Ethers.js", "封装账户、合约和交易回执读取"],
    ], [1.1, 1.5, 3.6])

    add_chapter(doc, "第三章  系统需求分析")
    add_section_heading(doc, "3.1  角色分析")
    add_para(doc, "系统主要涉及管理员、选民和审查者三类角色。管理员负责确定投票主题、候选项、投票截止时间和白名单地址，并执行部署脚本。选民通过 MetaMask 钱包连接前端，在投票窗口内选择候选项并提交交易。审查者不需要拥有投票资格，也可以通过前端、部署记录和 Etherscan 链接验证合约地址、交易 Hash、事件日志和结果统计。")
    add_table(doc, ["角色", "职责", "主要操作"], [
        ["管理员", "初始化投票主题、候选项、截止时间和 Merkle Root", "运行白名单脚本、部署合约、导出前端配置"],
        ["选民", "在投票窗口内完成一次有效投票", "连接钱包、选择候选项、签名交易、查看结果"],
        ["审查者", "验证投票过程和结果", "查看部署记录、交易 Hash、链上结果和审计报告"],
    ], [1.1, 2.5, 2.3])
    add_section_heading(doc, "3.2  功能需求")
    add_para(doc, "系统功能需求包括投票初始化、白名单生成、钱包连接、资格校验、投票提交、重复投票限制、投票窗口控制、实时结果查询和部署证据展示。投票初始化由部署脚本完成，白名单由脚本生成 Merkle Root 和每个地址对应的 proof。前端读取当前账户后匹配 proof，并在提交投票时将候选项索引和 proof 一并发送给合约。")
    add_table(doc, ["编号", "需求", "实现方式"], [
        ["F1", "配置投票主题和候选项", "部署脚本构造 VotingSystem 合约参数"],
        ["F2", "限制白名单投票资格", "Merkle Root 上链，proof 由前端提交"],
        ["F3", "防止重复投票", "mapping(address => bool) 记录投票状态"],
        ["F4", "控制投票截止时间", "block.timestamp 与 votingEndTime 比较"],
        ["F5", "公开查询投票结果", "getResults 返回候选项和票数数组"],
        ["F6", "同步前端部署配置", "exportFrontendArtifact 导出 ABI 与合约地址"],
    ], [0.8, 2.0, 3.1])
    add_section_heading(doc, "3.3  非功能需求")
    add_para(doc, "非功能需求主要包括安全性、可追溯性、可测试性、可部署性和易用性。安全性要求合约拒绝无效 proof、重复投票、非法候选项和超时投票；可追溯性要求部署地址、交易 Hash、Gas 和白名单根可记录；可测试性要求通过自动化测试覆盖主要分支；可部署性要求支持 Sepolia 测试网部署；易用性要求前端采用中文界面，能清晰展示钱包状态、投票面板、结果图表和链上证据。")
    add_para(doc, "安全性需求是系统设计的优先级最高约束。投票系统一旦接受了无资格地址、重复地址或错误候选项，就会直接破坏计票结果可信度。因此，本文将关键安全判断全部放在智能合约中执行，前端只做辅助提示。即使前端代码被篡改或用户直接使用脚本调用合约，链上 require 条件仍然会拒绝不符合规则的交易。")
    add_para(doc, "可追溯性需求要求每一次关键结果都能够被复核。合约部署后会产生部署交易 Hash，投票成功后会产生投票交易 Hash，合约状态可通过 getResults 等读取函数查询，Gas 消耗可由交易回执和 gas reporter 输出确认。这些证据共同构成论文中的“真实可追溯数据链”，避免毕业设计只停留在静态代码展示层面。")
    add_para(doc, "易用性需求并不意味着降低安全性，而是要求用户在安全规则下能够理解系统状态。前端需要用中文展示钱包地址、网络名称、投票窗口、是否已投票、白名单 proof 和交易回执等信息。对于答辩演示而言，界面不仅要能完成操作，还要能解释系统为什么允许或拒绝某一笔投票。")
    add_section_heading(doc, "3.4  威胁模型与设计边界")
    add_para(doc, "本系统重点防范的风险包括非白名单地址投票、同一地址重复投票、候选项索引越界、投票截止后继续投票和部署配置与前端展示不一致。系统通过智能合约校验和前端配置校验共同降低这些风险。")
    add_para(doc, "需要说明的是，本文系统实现的是地址维度的资格控制和地址伪匿名，并不提供真实身份认证、强匿名投票、抗胁迫投票和大规模隐私保护能力。链上交易公开可查，如果地址与真实身份在线下发生关联，投票行为仍可能被追踪。因此，本文将系统定位为去中心化电子投票原型和教学演示系统。")
    add_section_heading(doc, "3.5  典型用例分析")
    add_para(doc, "管理员用例的目标是创建一次可验证的投票活动。管理员在部署前确定投票标题、候选项、白名单地址和投票时长，随后运行白名单脚本得到 Merkle Root，并通过部署脚本将参数写入合约构造函数。部署完成后，管理员不再能够修改合约中的候选项、截止时间和白名单根，这一约束保证了投票规则不会在活动过程中被随意调整。")
    add_para(doc, "选民用例的目标是在投票窗口内完成一次有效投票。选民进入前端后连接 MetaMask，前端读取当前账户地址并在白名单 JSON 中查找 proof。如果当前地址存在 proof，前端允许用户选择候选项并提交交易；如果地址不在白名单中，前端会给出资格提示，即使用户绕过前端直接调用合约，也会因为 Merkle Proof 校验失败而被拒绝。")
    add_para(doc, "审查者用例的目标是复核投票结果。审查者可以不连接钱包，只通过合约地址、部署交易 Hash、投票交易 Hash 和公开读取接口查看结果。由于合约代码、交易记录和状态读取均在测试网上公开，审查者不需要信任前端页面本身，也可以使用 Etherscan 或脚本独立验证计票结果。")
    add_table(doc, ["用例", "前置条件", "正常流程", "异常处理"], [
        ["部署投票", "候选项不少于 2 个，白名单非空", "生成白名单、部署合约、导出前端配置", "参数非法时脚本或构造函数拒绝"],
        ["提交投票", "钱包连接，地址在白名单内，投票未截止", "选择候选项、签名交易、等待回执", "非白名单、重复投票、超时投票均回滚"],
        ["查询结果", "合约已部署", "调用 getResults 或查看前端统计", "RPC 异常时展示错误提示"],
    ], [1.0, 1.7, 2.2, 1.7])
    add_section_heading(doc, "3.6  本章小结")
    add_para(doc, "本章从角色、功能、非功能、威胁边界和典型用例五个角度明确了系统需求。需求分析表明，本文系统的核心不在于替代真实政务选举，而在于通过智能合约验证投票规则公开执行、资格控制可校验、计票结果可追溯这一技术路径。后续总体设计和详细实现均围绕这些需求展开，重点保证白名单验证、重复投票限制、投票窗口控制和部署配置同步。")

    add_chapter(doc, "第四章  系统总体设计")
    add_section_heading(doc, "4.1  总体架构设计")
    add_para(doc, "系统采用前端 DApp、钱包签名、区块链 RPC 和链上智能合约协同的架构。前端负责展示界面、加载部署配置、匹配白名单 proof、调用合约读写方法；MetaMask 负责账户授权和交易签名；Sepolia RPC 负责读取链上状态和广播交易；VotingSystem 合约负责最终资格校验、状态写入和事件触发。")
    add_para(doc, "该架构将“用户交互”和“可信执行”明确分离。前端可以提升使用体验，但不承担最终信任；钱包负责私钥隔离和交易确认，但不决定投票规则；RPC 节点负责连接链上网络，但不能绕过合约状态机；智能合约是投票规则的最终执行位置。通过这种分层设计，即使前端页面出现展示错误，审查者仍可以通过合约地址直接读取链上结果。")
    add_para(doc, "系统没有引入传统后端数据库，原因是投票资格和投票结果已经由白名单文件、链上合约和交易日志承载。对于毕业设计原型而言，减少后端服务可以降低部署复杂度，也能够突出区块链应用的核心特征。当然，在真实生产系统中，仍可能需要后端服务承担身份审核、消息通知、运维监控和审计日志聚合等职责。")
    add_picture(doc, FIG / "system-architecture.png", "图4.1  系统总体架构图", width=5.9)
    add_section_heading(doc, "4.2  业务流程设计")
    add_para(doc, "业务流程从投票配置开始。管理员先确定投票主题、候选项和白名单地址，随后运行白名单生成脚本得到 Merkle Root 与 proof 列表。部署脚本将主题、候选项、截止时间和 Merkle Root 写入合约构造参数。选民进入前端后连接钱包，前端读取当前账户并匹配 proof，最后提交投票交易。合约完成校验后更新票数，前端通过 getResults 展示公开统计结果。")
    add_picture(doc, FIG / "business-process.png", "图4.2  业务流程图", width=5.9)
    add_section_heading(doc, "4.3  投票交易流程设计")
    add_para(doc, "投票交易属于链上写操作，必须由选民钱包签名。前端只负责构造 candidateIndex 和 merkleProof 参数，不能绕过合约校验。合约执行时先完成投票窗口、候选项索引、重复投票状态和 proof 有效性检查，再写入 voted 映射并增加候选项票数，最后触发 VoteCast 事件。")
    add_picture(doc, FIG / "voting-flow.png", "图4.3  投票交易执行流程图", width=5.9)
    add_section_heading(doc, "4.4  合约状态与接口设计")
    add_para(doc, "VotingSystem 合约保存投票主题 title、不可变投票截止时间 votingEndTime、不可变 merkleRoot、候选项数组 candidates、票数数组 voteCounts 和 voted 映射。候选项与票数按相同索引对应，避免额外映射带来的遍历成本。合约对外提供 vote、getCandidates、getResults、isVotingOpen、hasAddressVoted 和 candidateCount 等接口。")
    add_table(doc, ["接口", "类型", "说明"], [
        ["vote(uint256,bytes32[])", "写入", "提交候选项索引和 Merkle Proof，成功后更新票数"],
        ["getResults()", "读取", "一次性返回候选项名称和票数数组"],
        ["isVotingOpen()", "读取", "根据区块时间判断投票窗口是否开放"],
        ["hasAddressVoted(address)", "读取", "查询指定地址是否已经投票"],
        ["candidateCount()", "读取", "返回候选项数量"],
    ], [2.0, 1.0, 3.2])
    add_section_heading(doc, "4.5  前端模块设计")
    add_para(doc, "前端采用组件化设计，主要包括控制台首页、钱包连接面板、投票提交面板、实时结果图表、链上证据面板和资格安全校验面板。状态管理由 useWallet 和 useVotingContract 两个 Hook 完成，前者负责钱包连接和网络识别，后者负责读取部署配置、加载白名单、实例化合约、读取链上结果和提交交易。")
    add_section_heading(doc, "4.6  数据一致性设计")
    add_para(doc, "系统中的数据分为链上状态、链下配置和前端展示三类。链上状态包括合约地址、候选项、票数、投票截止时间、白名单根和已投票映射，是最终可信来源；链下配置包括部署记录、前端 deployment.json 和 whitelist.json，用于帮助前端定位合约并构造交易参数；前端展示包括候选项列表、统计图表和交易回执，仅作为用户交互层，不作为最终计票依据。")
    add_para(doc, "为了避免部署后前端仍读取旧合约的问题，项目设计了 exportFrontendArtifact 脚本。该脚本从部署记录中读取最新合约地址、ABI 和构造参数，再统一写入 apps/web/public/deployment.json。前端启动时只读取该文件，不在代码中硬编码合约地址。这样，当 Sepolia 合约重新部署时，只需要重新运行导出脚本并构建前端，即可保证界面与链上合约一致。")
    add_para(doc, "白名单一致性由两个层面保证：第一，生成脚本对地址进行 checksum 规范化，并检测重复地址，避免同一地址因大小写不同被重复写入；第二，测试用例会校验前端 whitelist.json 中的 Merkle Root 与部署配置中的 merkleRoot 是否一致。若二者不一致，前端即使能够连接合约，也无法为对应地址生成有效 proof，最终会被合约拒绝。")
    add_table(doc, ["数据项", "保存位置", "一致性约束"], [
        ["合约地址", "deployments/sepolia.json、deployment.json", "前端配置必须来自最新部署记录"],
        ["ABI", "artifacts 与 public/deployment.json", "导出脚本同步，前端按 ABI 调用方法"],
        ["Merkle Root", "合约状态、部署记录、白名单文件", "三处值必须一致，否则 proof 无效"],
        ["投票结果", "合约 voteCounts", "前端只读取展示，不本地计票"],
    ], [1.2, 2.2, 3.0])
    add_section_heading(doc, "4.7  本章小结")
    add_para(doc, "综上，系统总体设计的重点在于把容易产生争议的状态放到链上，把可以离线计算的证明材料放到链下，把面向用户的展示逻辑放到前端。链上状态负责最终可信性，链下脚本负责降低成本和提高部署效率，前端负责解释和呈现结果。三者之间通过部署记录和自动化测试建立一致性约束，从而使系统既能运行，也能被复核。")
    add_para(doc, "本章完成了系统从业务流程到技术结构的总体设计。后续实现阶段需要重点保证三点：第一，合约接口必须与前端 ABI 保持一致；第二，白名单生成脚本必须与合约叶子哈希算法保持一致；第三，部署记录必须与前端配置保持一致。只要这三类一致性得到保障，系统演示时出现的问题就可以被快速定位到链上状态、链下配置或前端展示中的某一层。")

    add_chapter(doc, "第五章  系统详细实现")
    add_section_heading(doc, "5.1  开发环境与工程结构")
    add_para(doc, "项目采用 Node.js、npm、Hardhat、Solidity、React、Vite 和 TypeScript 构建。合约代码位于 contracts 目录，部署和白名单脚本位于 scripts 目录，前端应用位于 apps/web 目录，测试用例位于 test 目录，部署记录和论文支撑材料分别保存在 deployments、docs 和 deliverables 目录。")
    add_table(doc, ["目录或文件", "作用"], [
        ["contracts/VotingSystem.sol", "投票智能合约实现"],
        ["scripts/generateWhitelist.ts", "生成白名单 Merkle Root 和 proof"],
        ["scripts/deploy.ts", "部署合约到本地或 Sepolia 网络"],
        ["scripts/exportFrontendArtifact.ts", "导出前端 ABI 和部署配置"],
        ["apps/web/src", "React 前端源码"],
        ["test", "合约、脚本和前端资源自动化测试"],
        ["deliverables/graduation", "毕业设计说明书、截图、图表和报告"],
    ], [2.1, 4.2])
    add_section_heading(doc, "5.2  智能合约实现")
    add_para(doc, "合约构造函数对标题、候选项数量、投票截止时间和 Merkle Root 进行一次性校验，避免部署出不可用的合约实例。votingEndTime 和 merkleRoot 被声明为 immutable，部署后不可修改，既符合投票规则固定的业务要求，也减少了不必要的存储写入风险。")
    add_mono(doc, "require(bytes(_title).length > 0, \"Title is required\");\nrequire(_candidates.length >= 2, \"Invalid candidate count\");\nrequire(_votingEndTime > block.timestamp, \"Voting end time must be in the future\");\nrequire(_merkleRoot != bytes32(0), \"Merkle root is required\");")
    add_para(doc, "vote 函数采用 Checks-Effects-Interactions 思路。函数先进行所有输入和状态校验，再修改 voted 和 voteCounts。合约不接收 ETH，不向外部地址转账，也不调用不可信外部合约，因此不存在典型资金重入路径。")
    add_mono(doc, "require(block.timestamp <= votingEndTime, \"Voting has ended\");\nrequire(candidateIndex < candidates.length, \"Invalid candidate index\");\nrequire(!voted[msg.sender], \"Address has already voted\");\nbytes32 leaf = keccak256(abi.encodePacked(msg.sender));\nrequire(MerkleProof.verify(merkleProof, merkleRoot, leaf), \"Address is not eligible\");\nvoted[msg.sender] = true;\nvoteCounts[candidateIndex] += 1;")
    add_para(doc, "合约中的 voted 映射用于记录每个地址是否已经投票。该映射以 address 为键，以 bool 为值，查询成本为常数级。投票成功后先将 voted[msg.sender] 置为 true，再增加候选项票数。由于后续没有外部调用，即使从形式上看不存在重入入口，仍然保持“先检查、再修改状态”的顺序，有利于形成规范的合约编写习惯。")
    add_para(doc, "候选项和票数采用两个数组保存，索引一一对应。这样设计的优点是 getResults 可以一次性返回候选项名称数组和票数数组，前端无需多次调用合约读取每一个候选项。对于小规模投票场景，数组读取成本可控，也便于在测试中直接断言每个候选项的票数变化。")
    add_para(doc, "合约没有设置管理员修改函数，原因是本文投票活动采用一次性部署模型。投票主题、候选项、截止时间和白名单根在构造函数中确定，部署后不可调整。这种设计牺牲了动态配置能力，但换来了规则稳定性。对于毕业设计演示而言，重新部署一个新合约比在旧合约中引入管理员权限更能体现“规则一旦上链便不应随意改变”的思想。")
    add_para(doc, "事件 VoteCast 用于记录投票成功行为。虽然本系统前端主要通过 getResults 读取当前票数，但事件日志仍然具有审计意义。审查者可以根据事件中的 voter 和 candidateIndex 复核每一笔成功投票，也可以将事件日志与交易 Hash、区块高度和时间戳结合，形成完整的执行轨迹。")
    add_section_heading(doc, "5.3  白名单与 Merkle Proof 实现")
    add_para(doc, "白名单脚本接收地址数组后，对每个地址执行 checksum 规范化并拒绝重复地址。叶子节点计算方式与合约保持一致，即 keccak256(abi.encodePacked(address))。脚本生成 Merkle Root、每个地址的 leaf 和 proof，并写入前端 public 目录，供 DApp 在浏览器中读取。")
    add_para(doc, "当前白名单包含 3 个地址，Merkle Root 为 0xb84d167dee14c531723adc7c8625c29224727496b7a19329eda9cbc6d15c4a21。合约只保存根哈希，不保存完整地址列表，既减少链上存储，也避免直接在合约状态中暴露完整白名单。")
    add_section_heading(doc, "5.4  前端 DApp 实现")
    add_para(doc, "前端首页采用较深的棕黑色控制台风格，突出“去中心化电子投票系统”的毕业设计主题。页面顶部展示链上规则，主体区域展示系统说明、合约执行记录、钱包连接、投票提交、实时结果、链上证据和白名单校验状态。前端所有英文控制台名称和主要提示已替换为中文，使其更适合毕业设计演示。")
    add_para(doc, "钱包连接模块会读取当前账户地址和网络链 ID，并将地址缩写展示在页面中。若用户连接的不是 Sepolia 网络，前端会提示当前网络状态，避免将网络错误误判为合约调用错误。对于只想查看结果的审查者，页面仍然可以通过只读 RPC 加载部署配置和链上结果，不强制要求连接钱包。")
    add_para(doc, "投票面板根据候选项数组动态渲染选项。用户选择候选项后，前端调用合约 vote 方法并传入候选项索引和当前地址对应的 Merkle Proof。交易提交后，页面进入等待确认状态；交易确认后，前端刷新投票结果并展示交易 Hash。该流程使用户能够看到从选择候选项到链上确认的完整闭环。")
    add_para(doc, "结果统计模块将合约返回的候选项和票数转换为可视化柱状图，同时保留原始数值。对于投票系统而言，图表用于提高可读性，但不能替代链上数据。因此，页面同时展示合约地址、部署交易 Hash、白名单根和 ABI 摘要，便于答辩时说明图表数据来自合约读取，而不是前端本地模拟。")
    add_para(doc, "Merkle 安全面板用于解释白名单验证过程。页面展示当前账户是否命中白名单、proof 深度、叶子哈希和根哈希摘要。该设计面向毕业设计答辩场景，可以把原本隐藏在脚本和合约中的安全机制可视化，使评审能够理解为什么某个地址具备或不具备投票资格。")
    add_picture(doc, SHOT / "system-overview-top.png", "图5.1  系统首页与合约执行记录", width=5.9)
    add_picture(doc, SHOT / "wallet-and-voting-panel.png", "图5.2  钱包连接与投票提交面板", width=5.9)
    add_picture(doc, SHOT / "results-and-evidence-panel.png", "图5.3  实时投票结果与链上证据面板", width=5.9)
    add_picture(doc, SHOT / "merkle-security-panel.png", "图5.4  白名单与 Merkle Proof 安全校验面板", width=5.9)
    add_section_heading(doc, "5.5  部署与前端配置同步")
    add_para(doc, "部署流程由 npm 脚本串联完成。首先运行白名单生成脚本，得到最新 Merkle Root；然后运行 Hardhat 部署脚本，将合约部署到 Sepolia；最后运行前端导出脚本，把合约地址、ABI、候选项、截止时间和部署交易 Hash 写入 apps/web/public/deployment.json。这样前端读取的配置与链上部署保持一致。")
    add_mono(doc, "npm run generate:whitelist\nnpm run deploy:sepolia\nnpm run export:frontend -- --network sepolia\nnpm run web:build")
    add_section_heading(doc, "5.6  异常处理与交互细节")
    add_para(doc, "合约层异常以 require 条件表达式实现，主要覆盖标题为空、候选项数量不足、截止时间不在未来、Merkle Root 为空、候选项索引越界、重复投票、非白名单地址和投票已结束等情况。由于这些条件都在状态写入前执行，异常发生时交易会整体回滚，不会留下部分更新的票数或投票状态。")
    add_para(doc, "前端层异常处理主要面向用户操作体验。钱包未安装时，页面提示用户安装或启用 MetaMask；网络不是 Sepolia 时，页面显示当前网络状态，避免用户误以为本地配置错误；交易发送后，页面展示交易等待状态，交易确认后再刷新链上结果；如果 RPC 请求失败或用户拒绝签名，前端会将错误转换为中文提示，减少原始错误信息对非专业用户造成的理解负担。")
    add_para(doc, "白名单证明的读取采用前端本地匹配方式。当前账户地址变化时，前端重新读取 whitelist.json 并查找对应 proof；若匹配成功，则在投票提交时携带 proof；若匹配失败，前端仍显示系统状态和当前合约结果，但不会把用户误导为具备投票资格。该设计使界面能够同时服务选民和审查者两类角色。")
    add_para(doc, "在展示层面，前端将合约地址、部署交易 Hash、投票交易 Hash、白名单根、投票截止时间和 ABI 摘要放在同一控制台界面中。这样做的目的不是增加视觉复杂度，而是让答辩评审可以直接看到系统确实连接到测试网合约，并且能根据页面信息追踪到对应的链上记录。")

    add_chapter(doc, "第六章  测试、安全审计与部署分析")
    add_section_heading(doc, "6.1  测试环境与测试策略")
    add_para(doc, "系统测试采用自动化测试与链上实测相结合的方式。自动化测试覆盖前端显示工具、前端资源校验、部署脚本输入校验和 VotingSystem 合约核心逻辑。链上实测使用 Sepolia 测试网部署合约，并通过白名单地址提交真实投票交易，以验证部署配置、交易签名、链上状态读取和前端展示是否一致。")
    add_section_heading(doc, "6.2  自动化测试结果")
    add_para(doc, "运行 npm test 后，测试套件共 24 个用例全部通过。其中 VotingSystem 合约测试覆盖初始化、白名单投票、非白名单拒绝、重复投票拒绝、非法候选项拒绝、投票截止拒绝和结果读取等关键路径。")
    add_para(doc, "合约测试重点验证“允许路径”和“拒绝路径”是否同时成立。允许路径用于证明白名单地址在投票窗口内能够正常投票，并且票数会增加；拒绝路径用于证明非白名单地址、重复投票地址、非法候选项索引和超时交易均无法改变链上状态。只有这两类路径同时通过，才能说明合约不仅能完成正常业务，也能抵抗常见异常输入。")
    add_para(doc, "部署脚本测试主要验证工程入口的健壮性。例如候选项数量不足时应拒绝部署，白名单地址格式错误时应停止生成 proof，重复地址应被提前发现。这些测试看似不直接属于合约逻辑，但能够避免错误配置进入链上部署阶段，降低真实部署时浪费 Gas 或生成不可用合约的概率。")
    add_para(doc, "前端资源校验测试用于保证 deployment.json 和 whitelist.json 的一致性。由于 DApp 前端需要同时读取合约地址、ABI、Merkle Root 和 proof，如果这些文件由不同脚本或不同时间生成，就可能出现前端展示新合约、proof 却对应旧白名单的情况。通过自动化测试检查资源一致性，可以把这类问题提前暴露在构建阶段。")
    add_table(doc, ["测试模块", "用例数", "覆盖内容", "结果"], [
        ["前端显示工具", "7", "地址缩写、链 ID、票数统计、ABI 摘要、错误提示", "通过"],
        ["前端资源校验", "4", "部署配置、白名单、重复地址、Merkle Root 一致性", "通过"],
        ["部署脚本校验", "4", "候选项、地址格式、白名单生成与 proof 非空", "通过"],
        ["VotingSystem 合约", "9", "初始化、投票、拒绝非白名单、重复投票、截止时间、结果读取", "通过"],
    ], [1.5, 0.8, 3.2, 0.8])
    add_picture(doc, FIG / "test-result-table.png", "图6.1  测试结果汇总图", width=5.7)
    add_section_heading(doc, "6.3  Gas 消耗分析")
    add_para(doc, "Gas 分析使用 npm run test:gas 执行。该命令将 REPORT_GAS 设置为 true，并在 Hardhat 测试完成后输出合约部署和方法调用的 Gas 消耗。结果显示，本地部署 VotingSystem 消耗 767,754 Gas，vote 方法最小值为 73,654，最大值为 73,666，平均值为 73,661。真实 Sepolia 部署交易消耗 903,205 Gas，真实投票交易消耗 74,413 Gas。")
    add_para(doc, "vote 方法的 Gas 主要由四部分组成：一是读取合约状态并执行 require 校验；二是对 msg.sender 进行 Keccak-256 哈希并验证 Merkle Proof；三是首次写入 voted 映射；四是更新 voteCounts 数组。由于以太坊中写入存储槽的成本显著高于读取和内存计算，因此 voted 和 voteCounts 的状态更新是投票交易 Gas 的主要来源。")
    add_para(doc, "本地 gas reporter 输出的平均值和 Sepolia 投票回执中的真实值接近，说明测试环境对单次 vote 调用的估算具有参考价值。二者存在少量差异，可能来自编译优化、网络执行环境、交易封装和实际区块状态等因素。论文中同时保留本地测试数据和真实测试网回执，可以避免只依赖模拟结果。")
    add_para(doc, "如果后续要降低 Gas 成本，可以从减少存储写入、压缩事件参数和优化白名单证明长度等方向入手。但在本文系统中，安全性和可读性优先级高于极限优化。合约逻辑保持清晰，便于审计和答辩说明；Gas 消耗处于小规模投票可接受范围内，因此未进行过度压缩。")
    add_picture(doc, FIG / "gas-comparison-chart.png", "图6.2  Gas 消耗分析图", width=5.7)
    add_table(doc, ["操作", "Gas Used", "数据来源", "说明"], [
        ["本地部署", "767,754", "npm run test:gas", "初始化合约字节码和状态"],
        ["vote 平均值", "73,661", "npm run test:gas", "白名单投票平均消耗"],
        ["Sepolia 部署", "903,205", "部署交易回执", "真实测试网部署消耗"],
        ["Sepolia 投票", "74,413", "投票交易回执", "真实测试网投票消耗"],
    ], [1.4, 1.0, 1.4, 2.5])
    add_section_heading(doc, "6.4  安全审计与风险分析")
    add_para(doc, "本文使用 Slither 对合约进行静态分析，命令为 slither . --filter-paths \"node_modules|artifacts|cache\"。Slither 共分析 3 个合约对象和 101 个检测器，最终输出 3 条 timestamp 相关提示，位置分别为构造函数、vote 函数和 isVotingOpen 函数。")
    add_para(doc, "timestamp 风险的核心在于 block.timestamp 由区块生产者给出，理论上可能存在小范围偏移，不适合用于随机数、抽奖或资金分配。本系统仅使用 block.timestamp 判断投票截止时间，不依赖秒级精度产生经济收益，因此该风险属于可解释的业务边界。系统同时不存在 ETH 托管、外部合约调用和不可信回调，重入风险较低。")
    add_para(doc, "除 timestamp 外，本文还从访问控制、输入校验、重入风险、整数溢出和数据一致性等角度进行人工复核。访问控制方面，合约没有管理员后门函数，部署后任何账户都不能修改候选项和白名单根；输入校验方面，候选项索引、投票时间、重复投票和 proof 有效性均在状态写入前检查；整数安全方面，Solidity 0.8.24 默认启用溢出检查，票数递增不会静默回绕。")
    add_para(doc, "重入风险方面，vote 函数不转账、不调用外部合约，也不使用低级 call，因此不存在典型的外部回调入口。即便如此，代码仍按 Checks-Effects-Interactions 顺序组织，使安全规范体现在实现结构中。对于毕业设计而言，能够解释为什么某类风险不存在，比简单写一句“无风险”更符合工程审计要求。")
    add_para(doc, "前端安全方面，需要认识到前端不能作为最终安全边界。用户可以修改浏览器代码、直接使用脚本调用合约，或者构造任意 candidateIndex 和 proof。因此，本文只把前端作为交互和提示层，所有影响投票结果的判断都由合约重新执行。该设计符合区块链 DApp 中“前端可替换、合约可信执行”的基本原则。")
    add_para(doc, "私钥管理是系统实践中的重要风险。项目部署使用测试网账户完成，私钥只应保存在本地 .env 文件中，不应提交到 GitHub。由于测试网私钥在调试过程中存在暴露可能，完成毕业设计材料后应更换部署钱包或放弃该测试账户，避免后续被他人使用。")
    add_picture(doc, FIG / "security-audit-summary.png", "图6.3  Slither 安全审计结果摘要", width=5.7)
    add_section_heading(doc, "6.5  Sepolia 测试网部署结果")
    add_para(doc, "最新 Sepolia 合约地址为 0x506db4C4b9A63d127b3613AC87B2Cae3Fe185d2b，部署交易 Hash 为 0x25e62cd381f298afd1fb1ed0676c3126658663dd8304f56bd434ded272dbfcb8。部署账户为 0x372ee50901D62F3b314936C9302b19F8F477716E，投票窗口约为 3 天，截止时间为北京时间 2026 年 5 月 27 日 11:14:02。")
    add_para(doc, "重新部署合约的目的，是为公网部署和远程演示提供更充足的投票时间。与 30 分钟演示窗口相比，3 天窗口更适合让他人通过服务器访问前端并连接 Sepolia 钱包进行测试。部署完成后，前端配置文件被重新导出，确保页面中的合约地址与 Sepolia 上的最新实例一致。")
    add_para(doc, "部署证据不仅包括合约地址，还包括部署交易 Hash、部署账户、白名单根、Gas 消耗和前端读取截图。合约地址能够定位链上合约实例，交易 Hash 能够定位具体执行过程，Gas 消耗能够说明执行成本，截图能够证明前端已读取对应合约状态。历史投票交易回执继续保留为既有链上证据，但不作为本次 3 天新合约的投票结果。")
    add_picture(doc, FIG / "deployment-evidence-table.png", "图6.4  Sepolia 部署证据表", width=5.9)
    add_picture(doc, SHOT / "system-home-chain-read.png", "图6.5  前端读取真实链上投票结果截图", width=4.8)
    add_section_heading(doc, "6.6  可复现性与结果评价")
    add_para(doc, "为保证论文数据可复现，本文将关键命令输出保存为文本报告，包括 compile-output.txt、test-output.txt、gas-output.txt、slither-output.txt 和 web-build-output.txt。说明书中引用的测试通过数量、Gas 消耗、合约地址、部署交易 Hash 和投票交易 Hash 均来自这些报告或链上回执，而不是手工估算。")
    add_para(doc, "从测试覆盖结果看，系统已经覆盖了核心业务分支，能够证明合约在正常投票和异常输入下符合预期。从 Gas 结果看，vote 方法平均约 7.3 万 Gas，在小规模投票场景中成本可接受；部署交易 Gas 高于本地估算，主要与测试网真实交易环境和构造参数写入有关，但部署属于一次性操作，对普通选民无直接影响。")
    add_para(doc, "从安全审计结果看，Slither 未发现重入、未检查外部调用、任意转账、未初始化存储指针等高风险问题。timestamp 提示需要在论文中说明其适用范围：本系统用时间戳判断投票窗口，而不是生成随机数或分配资金，因此矿工或验证者的小范围时间偏移不会改变白名单资格、候选项索引和重复投票状态。")
    add_para(doc, "从系统演示效果看，前端能够读取真实测试网合约状态，并通过截图展示钱包状态、投票面板、结果统计和 Merkle Proof 信息。该结果说明系统已经形成从链下配置、合约部署、钱包签名、链上执行到前端展示的闭环，满足毕业设计“可执行程序”和“理论与实践结合”的基本要求。")

    add_chapter(doc, "第七章  总结与展望")
    add_section_heading(doc, "7.1  工作总结")
    add_para(doc, "本文围绕传统中心化电子投票系统存在的单点故障、数据易被人为干预和结果复核成本较高等问题，设计并实现了一个基于 Solidity 智能合约与 Ethereum Sepolia 测试网的去中心化电子投票系统。系统完成了从需求分析、总体设计、合约开发、前端实现、测试验证、安全审计到测试网部署的完整工程流程。")
    add_para(doc, "在技术实现上，系统使用 VotingSystem 合约固定投票规则，使用 Merkle Tree 白名单实现地址级资格控制，使用 mapping 防止重复投票，使用 Ethers.js 和 MetaMask 完成钱包交互，使用 Hardhat 组织测试与部署，使用 Slither 输出安全审计证据。最终系统能够在前端展示真实链上结果，并通过合约地址、交易 Hash、Gas 数据和测试报告进行追溯。")
    add_para(doc, "与仅完成一个前端页面或一个单独合约示例相比，本文更强调完整工程闭环。项目不仅包含智能合约代码，还包含白名单生成脚本、部署脚本、前端配置导出、自动化测试、Gas 报告、安全审计输出、测试网交易证据、运行截图和说明书材料。各部分之间通过真实数据连接，能够支撑毕业设计从“能运行”到“可解释、可复核”的要求。")
    add_para(doc, "从理论与实践结合角度看，本文在实现投票功能的同时，对 ECDSA 签名、Keccak-256 哈希、Merkle Proof 验证、哈希碰撞概率和 Gas 成本来源进行了分析。这样可以说明系统为什么能够防止伪造身份、为什么白名单根可以代表地址集合、为什么哈希碰撞在工程上可忽略，以及为什么链上状态写入会成为主要成本。")
    add_section_heading(doc, "7.2  不足与展望")
    add_para(doc, "受毕业设计周期和测试网场景限制，本文系统仍存在若干不足。第一，系统采用地址白名单，不包含真实身份认证和实名授权流程；第二，链上交易公开可查，无法提供强匿名和抗关联分析能力；第三，当前候选项和截止时间在部署后固定，不支持治理式动态变更；第四，系统面向小规模投票演示，尚未处理大规模并发、隐私保护和生产级运维监控问题。")
    add_para(doc, "后续可从四个方向继续完善：一是引入链下身份认证与链上凭证结合机制，提高资格发放可信度；二是研究零知识证明、盲签名或承诺揭示机制，增强投票隐私；三是扩展后台管理与多轮投票能力，提高系统可配置性；四是补充形式化验证、第三方审计和持续集成流程，使系统更接近生产级工程要求。")
    add_section_heading(doc, "7.3  工程伦理与合规说明")
    add_para(doc, "电子投票系统天然涉及公平性、隐私性和结果可信度等问题，因此在毕业设计中必须明确系统用途和适用边界。本文系统仅用于课程研究、技术验证和答辩演示，不面向真实公共选举、商业投票或具有法律效力的决策场景。系统中的白名单地址为测试地址，使用 Sepolia 测试币完成交易，不涉及真实资金托管和真实身份数据处理。")
    add_para(doc, "在数据保护方面，系统没有采集姓名、身份证号、手机号等敏感个人信息，只以 Ethereum 地址作为测试投票资格标识。需要注意的是，区块链地址并不等同于完全匿名身份，一旦地址与现实身份在其他场景中发生绑定，链上投票行为可能被关联分析。因此，若系统未来用于真实业务，应增加隐私保护、授权告知、数据最小化和合规审查机制。")
    add_para(doc, "在安全责任方面，本文保留了测试记录和审计结果，并对 Slither 提示进行解释，避免只展示成功路径而忽略风险边界。毕业设计交付材料中的合约地址、交易 Hash 和 Gas 数据均可追溯，能够接受复核。对于已经在开发过程中使用过的测试网私钥，应在项目提交后更换或弃用，避免因密钥泄露造成后续测试资产损失。")

    add_chapter(doc, "参考文献")
    refs = [
        "[1] Nakamoto S. Bitcoin: A Peer-to-Peer Electronic Cash System[EB/OL]. 2008.",
        "[2] Buterin V. Ethereum Whitepaper: A Next-Generation Smart Contract and Decentralized Application Platform[EB/OL]. 2014.",
        "[3] Wood G. Ethereum: A Secure Decentralised Generalised Transaction Ledger[R]. Ethereum Yellow Paper, 2024.",
        "[4] OpenZeppelin. Contracts Documentation: MerkleProof[EB/OL].",
        "[5] Solidity Team. Solidity Documentation, Version 0.8.x[EB/OL].",
        "[6] Hardhat. Ethereum Development Environment Documentation[EB/OL].",
        "[7] Trail of Bits. Slither: Static Analyzer for Solidity[EB/OL].",
        "[8] Johnson D, Menezes A, Vanstone S. The Elliptic Curve Digital Signature Algorithm (ECDSA)[J]. International Journal of Information Security, 2001.",
        "[9] Merkle R C. A Digital Signature Based on a Conventional Encryption Function[C]. CRYPTO, 1987.",
        "[10] Ethers.js Contributors. Ethers.js Documentation[EB/OL].",
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False)

    add_chapter(doc, "附录A  关键命令与运行证据")
    add_section_heading(doc, "A.1  主要命令")
    add_mono(doc, "npm install\nnpm run compile\nnpm test\nnpm run test:gas\nnpm run generate:whitelist\nnpm run deploy:sepolia\nnpm run export:frontend -- --network sepolia\nnpm run web:build\nslither . --filter-paths \"node_modules|artifacts|cache\"")
    add_section_heading(doc, "A.2  最新部署与交易证据")
    add_table(doc, ["字段", "内容"], [
        ["合约地址", "0x506db4C4b9A63d127b3613AC87B2Cae3Fe185d2b"],
        ["部署交易 Hash", "0x25e62cd381f298afd1fb1ed0676c3126658663dd8304f56bd434ded272dbfcb8"],
        ["历史投票交易 Hash", "0x9bc1ed5d902a984f86215e2ce01ffd0e3b95b041a101b5edc5fadd9b6adf629e"],
        ["白名单根", "0xb84d167dee14c531723adc7c8625c29224727496b7a19329eda9cbc6d15c4a21"],
        ["测试结果", "24 passing"],
        ["Slither 审计", "3 条 timestamp 提示，均已解释"],
    ], [1.4, 4.7])

    add_chapter(doc, "致谢")
    add_para(doc, "本毕业设计从选题、需求分析、系统设计到实现验证，得到了指导教师王立鹏老师的指导和帮助。老师在研究方向、系统边界、论文结构和工程规范方面提出了许多重要建议，使本文能够从单纯代码实现逐步完善为包含需求、设计、测试、部署和安全分析的完整毕业设计。")
    add_para(doc, "同时感谢学院在毕业设计过程中提供的模板、规范和阶段性检查要求。通过本课题的完成，作者进一步理解了区块链应用开发、智能合约安全、前端 DApp 交互和软件工程文档化的重要性。由于个人能力和时间有限，系统仍有不足之处，恳请各位老师批评指正。")


def main() -> None:
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    # The school template is used as the formatting and structure reference. A
    # clean DOCX package is generated here because the original template carries
    # legacy OLE/media relationships that make LibreOffice PDF rendering fail.
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    start_declaration_section(doc)
    add_declarations(doc)
    start_main_section(doc)
    add_abstracts(doc)
    add_toc(doc)
    add_body(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
